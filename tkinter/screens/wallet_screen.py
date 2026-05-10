import sys, os as _os
sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))

import tkinter as tk
from tkinter import ttk
from constants import *
from datetime import datetime
from PIL import Image, ImageTk
from db import (get_wallets, get_wallet_transactions,
                get_wallet_receipts, get_wallet_reports, get_wallet_budget,
                add_wallet_transaction, update_wallet_transaction,
                delete_wallet_transaction, upsert_wallet_budget,
                add_financial_report, get_latest_report,
                add_wallet_receipt, get_receipt_public_url,
                get_receipt_download_url, delete_receipt,
                download_receipt_bytes, submit_report,
                get_wallet_archives)
import db as _db_module

def _sb_client():
    """Return the shared Supabase client from db.py."""
    return _db_module._sb

_ASSETS      = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), "assets", "images")
_CARD_COLORS = ["#F3D58D","#D4E8C2","#C2D4E8","#E8C2D4","#D4C2E8","#C2E8D4","#E8D4C2"]
_INCOME_CLR  = "#2E7D32"
_EXPENSE_CLR = "#C62828"
_ACTIVE_TAB  = "#E59E2C"
_FILTER_ACT  = "#E59E2C"
_FILTER_BDR  = "#ECDDC6"
_BTN_BROWN   = "#A24A00"
_BTN_HOV     = "#8B3A00"


def _img(name, w, h, cache):
    path = _os.path.join(_ASSETS, name)
    if not _os.path.exists(path):
        return None
    try:
        ph = ImageTk.PhotoImage(Image.open(path).resize((w, h), Image.LANCZOS))
        cache.append(ph)
        return ph
    except Exception:
        return None


def _img_rounded(name, w, h, radius, cache):
    """Load image with rounded corners using a PIL mask."""
    path = _os.path.join(_ASSETS, name)
    if not _os.path.exists(path):
        return None
    try:
        from PIL import ImageDraw
        img = Image.open(path).resize((w, h), Image.LANCZOS).convert("RGBA")
        mask = Image.new("L", (w, h), 0)
        draw = ImageDraw.Draw(mask)
        draw.rounded_rectangle((0, 0, w, h), radius=radius, fill=255)
        img.putalpha(mask)
        ph = ImageTk.PhotoImage(img)
        cache.append(ph)
        return ph
    except Exception:
        return _img(name, w, h, cache)


class WalletScreen(tk.Frame):
    def __init__(self, parent, org=None, **kwargs):
        super().__init__(parent, bg=BG_CREAM, **kwargs)
        self._org   = org or {}
        self._imgs  = []
        self._view  = "list"          # "list" | "detail"
        self._wallet = None           # selected wallet dict
        self._build()

    # ── Shared pill-button factory (same smooth rounding as sidebar) ──
    def _make_pill_btn(self, parent, text, bg_hex, hover_hex,
                       fg="white", width=None, height=36,
                       font_spec=None, side=None, padx=0, pady=0,
                       command=None):
        """
        Returns a tk.Canvas that looks like a fully-rounded pill button,
        drawn with PIL at 4× scale for smooth anti-aliased edges —
        identical technique to the sidebar nav pills.
        """
        from PIL import Image, ImageDraw as _ID

        fnt = font_spec or font(9, "bold")
        cv = tk.Canvas(parent, bd=0, highlightthickness=0,
                       bg=parent.cget("bg"), cursor="hand2",
                       height=height)
        if width:
            cv.config(width=width)
        if side is not None:
            cv.pack(side=side, padx=padx, pady=pady)
        else:
            cv.pack(padx=padx, pady=pady)

        def _draw(hover=False):
            cv.delete("all")
            bw = cv.winfo_width() or (width or 80)
            bh = cv.winfo_height() or height
            if bw < 4 or bh < 4:
                return
            # use _active_bg if set (for filter button active state), else normal
            active_bg = getattr(cv, "_active_bg", None)
            active_fg = getattr(cv, "_active_fg", None)
            is_active = active_bg is not None
            colour    = active_bg if is_active else (hover_hex if hover else bg_hex)
            txt_colour = active_fg if is_active else fg
            scale = 4
            sw, sh = bw * scale, bh * scale
            r = (sh // 2)          # fully rounded — same as sidebar
            img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
            cr = int(colour.lstrip("#"), 16)
            rgb = ((cr >> 16) & 255, (cr >> 8) & 255, cr & 255)
            # inactive (not active, not hover): draw outline border only
            if not is_active and not hover:
                # border colour from _FILTER_BDR (#ECDDC6)
                bdr_hex = _FILTER_BDR.lstrip("#")
                bdr_rgb = (int(bdr_hex[0:2],16), int(bdr_hex[2:4],16), int(bdr_hex[4:6],16))
                _ID.Draw(img).rounded_rectangle(
                    [0, 0, sw - 1, sh - 1], radius=r,
                    fill=rgb + (255,),
                    outline=bdr_rgb + (255,), width=6)
            else:
                _ID.Draw(img).rounded_rectangle(
                    [0, 0, sw - 1, sh - 1], radius=r,
                    fill=rgb + (255,))
            img = img.resize((bw, bh), Image.LANCZOS)
            ph = ImageTk.PhotoImage(img)
            cv._ph = ph
            cv.create_image(0, 0, anchor="nw", image=ph)
            cv.create_text(bw // 2, bh // 2, text=text,
                           fill=txt_colour, font=fnt, anchor="center")

        cv._redraw = _draw   # expose so _update_flt_styles can call it
        cv.bind("<Configure>", lambda e: _draw())
        cv.bind("<Enter>",     lambda e: _draw(hover=True))
        cv.bind("<Leave>",     lambda e: _draw(hover=False))
        if command:
            cv.bind("<Button-1>", lambda e: command())
        return cv

    # ── outer scaffold ────────────────────────────────────────────────
    def _build(self):
        outer = tk.Frame(self, bg=BG_CREAM, padx=12, pady=16)
        outer.pack(fill="both", expand=True)

        box_canvas = tk.Canvas(outer, bg=BG_CREAM, bd=0, highlightthickness=0)
        box_canvas.pack(fill="both", expand=True)

        self._box = tk.Frame(box_canvas, bg=BG_WHITE, padx=30, pady=24)
        box_win = box_canvas.create_window(0, 0, anchor="nw", window=self._box)

        def _draw_box_bg(event=None):
            w = box_canvas.winfo_width()
            h = box_canvas.winfo_height()
            if w < 2 or h < 2:
                return
            r = 20
            box_canvas.itemconfig(box_win, width=w - r * 2, height=h - r * 2)
            box_canvas.coords(box_win, r, r)
            scale = 4
            sw, sh = w * scale, h * scale
            from PIL import Image, ImageDraw
            cr = int(BG_CREAM.lstrip("#"), 16)
            bg_rgb = ((cr >> 16) & 255, (cr >> 8) & 255, cr & 255)
            img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
            img_bg = Image.new("RGBA", (sw, sh), bg_rgb + (255,))
            ImageDraw.Draw(img).rounded_rectangle(
                [0, 0, sw - 1, sh - 1], radius=r * scale,
                fill=(255, 255, 255, 255))
            img_bg.paste(img, mask=img)
            img_bg = img_bg.resize((w, h), Image.LANCZOS)
            ph = ImageTk.PhotoImage(img_bg)
            box_canvas._bg_ph = ph
            box_canvas.delete("box_bg")
            box_canvas.create_image(0, 0, anchor="nw", image=ph, tags="box_bg")
            box_canvas.tag_lower("box_bg")

        box_canvas.bind("<Configure>", _draw_box_bg)
        self._show_list_view()

    # ══════════════════════════════════════════════════════════════════
    # LIST VIEW
    # ══════════════════════════════════════════════════════════════════
    def _show_list_view(self):
        self._clear_box()
        self._view = "list"

        # ── header row ────────────────────────────────────────────────
        hdr = tk.Frame(self._box, bg=BG_WHITE)
        hdr.pack(fill="x", pady=(0, 10))
        tk.Label(hdr, text="Wallets", bg=BG_WHITE, fg=TEXT_DARK,
                 font=("Georgia", 22, "italic")).pack(side="left")

        # academic year dropdown
        self._year_var = tk.StringVar(value="All years")
        self._year_menu = ttk.Combobox(hdr, textvariable=self._year_var,
                                       state="readonly", width=12,
                                       font=font(9))
        self._year_menu.pack(side="right", padx=(6, 0))

        # search bar
        self._search_var = tk.StringVar()
        search = tk.Entry(hdr, textvariable=self._search_var,
                          font=font(9), bd=1, relief="solid",
                          highlightbackground=_FILTER_BDR,
                          highlightthickness=1)
        search.insert(0, "Search wallet...")
        search.config(fg=TEXT_MUTED)
        search.pack(side="right", ipady=5, ipadx=8)

        def _focus_in(e):
            if search.get() == "Search wallet...":
                search.delete(0, "end")
                search.config(fg=TEXT_DARK)
        def _focus_out(e):
            if not search.get():
                search.insert(0, "Search wallet...")
                search.config(fg=TEXT_MUTED)
        search.bind("<FocusIn>",  _focus_in)
        search.bind("<FocusOut>", _focus_out)

        # loading
        loading = tk.Label(self._box, text="Loading wallets...",
                           bg=BG_WHITE, fg=TEXT_MUTED, font=font(10))
        loading.pack(pady=30)
        self.update()

        try:
            all_wallets = get_wallets(self._org.get("id"))
        except Exception:
            all_wallets = []
        loading.destroy()

        if not all_wallets:
            self._empty(self._box, "No wallets found.",
                        "No active wallets for this organization.")
            return

        # build academic year list
        acad_years = sorted(set(w["acad_year"] for w in all_wallets), reverse=True)
        self._year_menu["values"] = ["All years"] + acad_years

        # default to current academic year
        now = datetime.now()
        cur_ay = (f"{now.year}-{now.year+1}" if now.month >= 8
                  else f"{now.year-1}-{now.year}")
        if cur_ay in acad_years:
            self._year_var.set(cur_ay)
        else:
            self._year_var.set(acad_years[0] if acad_years else "All years")

        # scrollable area
        canvas = tk.Canvas(self._box, bg=BG_WHITE, bd=0, highlightthickness=0)
        sb = ttk.Scrollbar(self._box, orient="vertical", command=canvas.yview)
        self._grid_inner = tk.Frame(canvas, bg=BG_WHITE)
        self._grid_inner.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        self._canvas_win = canvas.create_window((0, 0), window=self._grid_inner, anchor="nw")
        canvas.configure(yscrollcommand=sb.set)
        # keep inner frame same width as canvas
        canvas.bind("<Configure>",
            lambda e: canvas.itemconfig(self._canvas_win, width=e.width))
        canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self._bind_scroll(canvas)

        self._all_wallets  = all_wallets
        self._wallet_cards = []
        self._wallet_img   = None   # image is loaded dynamically per card at render time

        # wire up filters — rebuild grid on change
        self._search_var.trace_add("write", lambda *_: self._apply_filters())
        self._year_var.trace_add("write",   lambda *_: self._apply_filters())

        # initial render
        self._apply_filters()

    def _apply_filters(self):
        q  = self._search_var.get().lower()
        if q == "search wallet...":
            q = ""
        ay = self._year_var.get()

        # destroy all existing cards and rows
        for w in self._grid_inner.winfo_children():
            w.destroy()
        self._wallet_cards = []

        visible = [
            w for w in self._all_wallets
            if (q in w["name"].lower() or q in w["month"].lower())
            and (ay == "All years" or w["acad_year"] == ay)
        ]

        if not visible:
            tk.Label(self._grid_inner, text="No wallets match.",
                     bg=BG_WHITE, fg=TEXT_MUTED, font=font(9)).pack(pady=20)
            return

        cols = 4
        row_f = None
        for i, w in enumerate(visible):
            if i % cols == 0:
                row_f = tk.Frame(self._grid_inner, bg=BG_WHITE)
                row_f.pack(fill="x", pady=(0, 8))
                for c in range(cols):
                    row_f.columnconfigure(c, weight=1, uniform="walletcol")
            color = _CARD_COLORS[i % len(_CARD_COLORS)]
            card = self._wallet_card(row_f, w, color, self._wallet_img, col=i % cols)
            self._wallet_cards.append((card, w))

    def _wallet_card(self, parent, w, color, wallet_img, col=0):
        outer = tk.Frame(parent, bg="white", cursor="hand2")
        outer.grid(row=0, column=col, padx=(0, 8), pady=4, sticky="nsew")

        # canvas fills the outer frame — height is fixed, width is dynamic
        CARD_H = 185
        cv = tk.Canvas(outer, height=CARD_H,
                       bd=0, highlightthickness=0, bg="white")
        cv.pack(fill="x", expand=True)

        imgs_ref = self._imgs   # capture ref to screen-level image cache

        def _redraw(event=None):
            cw = cv.winfo_width()
            if cw < 10:
                return
            cv.delete("all")
            try:
                from PIL import Image as _Img, ImageDraw as _ID, ImageTk as _ITk
                raw = _Img.open(
                    _os.path.join(_ASSETS, "wallet.png")
                ).resize((cw, CARD_H), _Img.LANCZOS).convert("RGBA")
                # apply small rounded corners (radius=10)
                mask = _Img.new("L", (cw, CARD_H), 0)
                _ID.Draw(mask).rounded_rectangle(
                    (0, 0, cw - 1, CARD_H - 1), radius=10, fill=255)
                bg = _Img.new("RGBA", (cw, CARD_H), (255, 255, 255, 255))
                bg.paste(raw, mask=mask)
                ph = _ITk.PhotoImage(bg.convert("RGB"))
                imgs_ref.append(ph)   # keep alive in screen-level cache
                cv.create_image(0, 0, image=ph, anchor="nw")
            except Exception:
                pass

            month_name = w.get("month_name", "").upper()
            pad_x, pad_y = 6, 3
            text_h = 16 + pad_y * 2
            x1 = 8
            y1 = CARD_H - 10 - text_h
            x2 = x1 + len(month_name) * 7 + pad_x * 2
            y2 = CARD_H - 10
            cv.create_rectangle(x1, y1, x2, y2, fill="white", outline="", width=0)
            cv.create_text(x1 + pad_x, (y1 + y2) // 2,
                           text=month_name, anchor="w",
                           font=font(8, "bold"), fill=TEXT_DARK)

        cv.bind("<Configure>", _redraw)

        def _enter(e): cv.config(bg="#F5F1E8")
        def _leave(e): cv.config(bg="white")
        def _click(e, wallet=w): self._show_detail_view(wallet)

        for widget in (outer, cv):
            widget.bind("<Button-1>", _click)
            widget.bind("<Enter>",    _enter)
            widget.bind("<Leave>",    _leave)

        outer._wallet_name = w["name"]
        return outer

    # ══════════════════════════════════════════════════════════════════
    # DETAIL VIEW
    # ══════════════════════════════════════════════════════════════════
    def _show_detail_view(self, wallet):
        self._clear_box()
        self._view      = "detail"
        self._wallet    = wallet
        self._tx_filter = "all"

        # ── header ────────────────────────────────────────────────────
        hdr = tk.Frame(self._box, bg=BG_WHITE)
        hdr.pack(fill="x", pady=(0, 4))

        # back
        back = tk.Label(hdr, text="‹", bg=BG_WHITE, fg=_BTN_BROWN,
                        font=font(22), cursor="hand2")
        back.pack(side="left")
        back.bind("<Button-1>", lambda e: self._show_list_view())
        back.bind("<Enter>",    lambda e: back.config(fg="#5A2D0C"))
        back.bind("<Leave>",    lambda e: back.config(fg=_BTN_BROWN))

        # title — month name only, italic Georgia
        month_title = wallet.get("month_name", "").capitalize()
        tk.Label(hdr, text=month_title, bg=BG_WHITE, fg=TEXT_DARK,
                 font=("Georgia", 16, "italic")).pack(side="left", padx=(8, 0))

        # —— Add budget secondary button (right side) — rounded pill ——
        try:
            _cur_budget = get_wallet_budget(wallet["id"], wallet.get("year"), wallet.get("month_id"))
        except Exception:
            _cur_budget = 0.0

        _budget_text_ref = [f"Budget: Php {_cur_budget:,.2f}" if _cur_budget else "Add budget for this month"]

        from PIL import Image, ImageDraw as _ID2

        budget_cv = tk.Canvas(hdr, bd=0, highlightthickness=0,
                              bg=BG_WHITE, cursor="hand2", height=30)
        budget_cv.pack(side="right", padx=(6, 0))

        def _draw_budget_btn(hover=False):
            budget_cv.delete("all")
            # measure text width to size the canvas dynamically
            import tkinter.font as tkfont
            fnt_obj = tkfont.Font(family="TkDefaultFont", size=8)
            txt_w = fnt_obj.measure(_budget_text_ref[0])
            bw = txt_w + 28          # 14px padding each side
            bh = 30
            budget_cv.config(width=bw)
            scale = 4
            sw, sh = bw * scale, bh * scale
            r = sh // 2              # fully rounded pill
            fill_hex = "#F5F1E8" if hover else BG_WHITE
            cr = int(fill_hex.lstrip("#"), 16)
            fill_rgb = ((cr >> 16) & 255, (cr >> 8) & 255, cr & 255)
            bdr_hex = _FILTER_BDR.lstrip("#")
            bdr_rgb = (int(bdr_hex[0:2],16), int(bdr_hex[2:4],16), int(bdr_hex[4:6],16))
            img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
            _ID2.Draw(img).rounded_rectangle(
                [0, 0, sw - 1, sh - 1], radius=r,
                fill=fill_rgb + (255,),
                outline=bdr_rgb + (255,), width=6)
            img = img.resize((bw, bh), Image.LANCZOS)
            ph = ImageTk.PhotoImage(img)
            budget_cv._ph = ph
            budget_cv.create_image(0, 0, anchor="nw", image=ph)
            budget_cv.create_text(bw // 2, bh // 2,
                                  text=_budget_text_ref[0],
                                  fill="#616161", font=font(8), anchor="center")

        budget_cv.bind("<Configure>", lambda e: _draw_budget_btn())
        budget_cv.bind("<Enter>",     lambda e: _draw_budget_btn(hover=True))
        budget_cv.bind("<Leave>",     lambda e: _draw_budget_btn(hover=False))
        budget_cv.bind("<Button-1>",  lambda e: self._open_budget_dialog(
            budget_cv, _budget_text_ref, _draw_budget_btn))
        budget_cv.after(10, _draw_budget_btn)

        # —— + Add button with dropdown ——
        add_wrap = tk.Frame(hdr, bg=BG_WHITE)
        add_wrap.pack(side="right", padx=(0, 6))

        add_btn = self._make_pill_btn(
            add_wrap, "+ Add",
            bg_hex=_BTN_BROWN, hover_hex=_BTN_HOV,
            fg="white", width=80, height=34,
            font_spec=font(9, "bold"))
        add_btn.pack()

        # dropdown
        drop = tk.Frame(self._box, bg="white",
                        highlightbackground="#ddd", highlightthickness=1)
        drop._open = False

        for txt, kind in [("Add income transaction", "income"),
                          ("Add expense transaction", "expense"),
                          ("Add receipt", "receipt")]:
            item = tk.Label(drop, text=txt, bg="white", fg=TEXT_DARK,
                            font=font(9), padx=16, pady=8, anchor="w",
                            cursor="hand2")
            item.pack(fill="x")
            item.bind("<Enter>", lambda e, b=item: b.config(bg="#F5F1E8"))
            item.bind("<Leave>", lambda e, b=item: b.config(bg="white"))
            if kind == "receipt":
                item.bind("<Button-1>",
                          lambda e: (self._close_drop(drop),
                                     self._open_receipt_dialog()))
            else:
                item.bind("<Button-1>",
                          lambda e, k=kind: (self._close_drop(drop),
                                             self._open_tx_dialog(k)))

        def _toggle(e=None):
            if drop._open:
                self._close_drop(drop)
            else:
                drop.place(in_=self._box, relx=1.0, rely=0,
                           anchor="ne", x=-36, y=60)
                drop.lift()
                drop._open = True

        add_btn.bind("<Button-1>", _toggle)

        # ── tab nav ───────────────────────────────────────────────────
        tab_bar = tk.Frame(self._box, bg=BG_WHITE,
                           highlightbackground=_FILTER_BDR,
                           highlightthickness=0)
        tab_bar.pack(fill="x", pady=(12, 0))

        # bottom border line
        tk.Frame(self._box, bg=_FILTER_BDR, height=2).pack(fill="x")

        self._tab_btns = {}
        self._tab_content = tk.Frame(self._box, bg=BG_WHITE)
        self._tab_content.pack(fill="both", expand=True, pady=(12, 0))

        for key, label in [("transactions","Transactions"),("reports","Reports"),
                            ("receipts","Receipts"),("archives","Archive")]:
            btn = tk.Label(tab_bar, text=label, bg=BG_WHITE, fg="#616161",
                           font=font(10, "bold"), padx=20, pady=10,
                           cursor="hand2")
            btn.pack(side="left")
            btn.bind("<Button-1>", lambda e, k=key: self._switch_tab(k))
            btn.bind("<Enter>",    lambda e, b=btn, k=key: b.config(fg=TEXT_DARK) if k != self._active_tab else None)
            btn.bind("<Leave>",    lambda e, b=btn, k=key: b.config(fg="#616161") if k != self._active_tab else None)
            self._tab_btns[key] = btn

        self._active_tab = "transactions"
        self._switch_tab("transactions")

    def _close_drop(self, drop):
        drop.place_forget()
        drop._open = False

    def _open_tx_dialog(self, kind, existing=None):
        """Modal overlay for Add/Edit Transaction."""
        wallet  = self._wallet
        is_edit = existing is not None

        root = self.winfo_toplevel()
        root.update_idletasks()
        x, y = root.winfo_rootx(), root.winfo_rooty()
        w, h = root.winfo_width(), root.winfo_height()
        try:
            from PIL import ImageGrab, ImageEnhance
            screenshot = ImageGrab.grab(bbox=(x, y, x + w, y + h))
            screenshot = ImageEnhance.Brightness(screenshot).enhance(0.5)
            _bg_photo = ImageTk.PhotoImage(screenshot)
        except Exception:
            _bg_photo = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_canvas = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_photo:
            ov_canvas.create_image(0, 0, anchor="nw", image=_bg_photo)
            ov_canvas._bg_photo = _bg_photo
        else:
            ov_canvas.config(bg="black")
            ov_canvas.create_rectangle(0, 0, w, h, fill="black", stipple="gray50", outline="")
        ov_canvas.bind("<Button-1>", lambda e: overlay.destroy())

        modal = tk.Frame(overlay, bg=BG_WHITE, padx=28, pady=24,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=420)
        modal.bind("<Button-1>", lambda e: "break")

        hdr = tk.Frame(modal, bg=BG_WHITE)
        hdr.pack(fill="x", pady=(0, 4))

        if is_edit:
            title_text = "Edit Income Transaction" if kind == "income" else "Edit Expense Transaction"
        else:
            title_text = "Add Income Transaction" if kind == "income" else "Add Expense Transaction"
        tk.Label(hdr, text=title_text, bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(12, "bold")).pack(side="left")

        close_btn = tk.Label(hdr, text="×", bg=BG_WHITE, fg="#616161",
                             font=("Arial", 16), cursor="hand2")
        close_btn.pack(side="right")
        close_btn.bind("<Button-1>", lambda e: overlay.destroy())

        if is_edit:
            subtitle = "Update this income transaction." if kind == "income" else "Update this expense transaction."
        else:
            subtitle = "Record an income transaction for this wallet." if kind == "income"                        else "Record an expense transaction for this wallet."
        tk.Label(modal, text=subtitle, bg=BG_WHITE, fg="#616161",
                 font=font(8)).pack(anchor="w", pady=(0, 14))

        err_labels = {}

        def _field(parent, label_text, field_key):
            grp = tk.Frame(parent, bg=BG_WHITE)
            grp.pack(fill="x", pady=(0, 10))
            tk.Label(grp, text=label_text, bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(8, "bold")).pack(anchor="w")
            return grp

        def _entry(grp, field_key, **kw):
            e = tk.Entry(grp, font=font(9), bd=1, relief="solid",
                         highlightbackground=_FILTER_BDR,
                         highlightthickness=1, **kw)
            e.pack(fill="x", ipady=6, pady=(3, 0))
            err = tk.Label(grp, text="", bg=BG_WHITE, fg="#C62828", font=font(7))
            err.pack(anchor="w")
            err_labels[field_key] = err
            return e

        def _combobox(grp, field_key, values):
            var = tk.StringVar()
            cb = ttk.Combobox(grp, textvariable=var, values=values,
                              state="readonly", font=font(9))
            cb.pack(fill="x", pady=(3, 0))
            err = tk.Label(grp, text="", bg=BG_WHITE, fg="#C62828", font=font(7))
            err.pack(anchor="w")
            err_labels[field_key] = err
            return var

        def _show_err(key, msg):
            if key in err_labels: err_labels[key].config(text=msg)

        def _clear_err(key):
            if key in err_labels: err_labels[key].config(text="")

        grp_date = _field(modal, "Date Issued", "date")
        e_date = _entry(grp_date, "date")
        e_date.insert(0, existing["date_issued"][:10] if is_edit else datetime.now().strftime("%Y-%m-%d"))

        grp_qty = _field(modal, "Quantity", "qty")
        e_qty = _entry(grp_qty, "qty")
        if is_edit: e_qty.insert(0, str(existing.get("quantity", "")))

        if kind == "income":
            grp_type = _field(modal, "Type of Income", "income_type")
            income_types = ["Income Generating Projects", "Registration Fee", "Membership Fee"]
            v_income_type = _combobox(grp_type, "income_type", income_types)
            if is_edit and existing.get("income_type"):
                v_income_type.set(existing["income_type"])
        else:
            grp_part = _field(modal, "Particulars", "particulars")
            e_particulars = _entry(grp_part, "particulars")
            if is_edit: e_particulars.insert(0, existing.get("particulars") or "")

        grp_desc = _field(modal, "Description", "desc")
        e_desc = _entry(grp_desc, "desc")
        if is_edit: e_desc.insert(0, existing.get("description") or "")

        grp_price = _field(modal, "Price", "price")
        e_price = _entry(grp_price, "price")
        if is_edit: e_price.insert(0, str(existing.get("price", "")))

        footer = tk.Frame(modal, bg=BG_WHITE)
        footer.pack(fill="x", pady=(14, 0))

        cancel = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                          font=font(9, "bold"), padx=18, pady=8, cursor="hand2",
                          highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel.pack(side="left")
        cancel.bind("<Button-1>", lambda e: overlay.destroy())
        cancel.bind("<Enter>", lambda e: cancel.config(bg="#F5F1E8"))
        cancel.bind("<Leave>", lambda e: cancel.config(bg=BG_WHITE))

        save_btn = tk.Label(footer, text="Save", bg=_BTN_BROWN, fg="white",
                            font=font(9, "bold"), padx=18, pady=8, cursor="hand2")
        save_btn.pack(side="right")
        save_btn.bind("<Enter>", lambda e: save_btn.config(bg=_BTN_HOV))
        save_btn.bind("<Leave>", lambda e: save_btn.config(bg=_BTN_BROWN))

        def _save(e=None):
            valid = True
            date_val  = e_date.get().strip()
            qty_val   = e_qty.get().strip()
            desc_val  = e_desc.get().strip()
            price_val = e_price.get().strip()

            if kind == "income":
                type_val = v_income_type.get().strip()
                part_val = ""
            else:
                type_val = ""
                part_val = e_particulars.get().strip()

            if not date_val:
                _show_err("date", "Date is required."); valid = False
            else:
                _clear_err("date")

            if not qty_val:
                _show_err("qty", "Quantity is required."); valid = False
            elif not qty_val.isdigit() or int(qty_val) < 1:
                _show_err("qty", "Quantity must be a positive number."); valid = False
            else:
                _clear_err("qty")

            if kind == "income":
                if not type_val:
                    _show_err("income_type", "Type of income is required."); valid = False
                else:
                    _clear_err("income_type")
            else:
                if not part_val:
                    _show_err("particulars", "Particulars are required."); valid = False
                else:
                    _clear_err("particulars")

            if not desc_val:
                _show_err("desc", "Description is required."); valid = False
            else:
                _clear_err("desc")

            if not price_val:
                _show_err("price", "Price is required."); valid = False
            else:
                try:
                    float(price_val)
                    _clear_err("price")
                except ValueError:
                    _show_err("price", "Price must be a valid number."); valid = False

            if not valid:
                return

            try:
                if is_edit:
                    update_wallet_transaction(
                        tx_id       = existing["id"],
                        date_issued = date_val,
                        quantity    = int(qty_val),
                        income_type = type_val if kind == "income" else None,
                        particulars = part_val if kind == "expense" else None,
                        description = desc_val,
                        price       = float(price_val),
                    )
                else:
                    add_wallet_transaction(
                        wallet_id   = wallet["id"],
                        budget_id   = wallet["budget_id"],
                        kind        = kind,
                        date_issued = date_val,
                        quantity    = int(qty_val),
                        income_type = type_val if kind == "income" else None,
                        particulars = part_val if kind == "expense" else None,
                        description = desc_val,
                        price       = float(price_val),
                    )
                overlay.destroy()
                if self._active_tab == "transactions":
                    self._load_transactions()
            except Exception as err:
                _show_err("price", f"Failed to save: {err}")

        save_btn.bind("<Button-1>", _save)

    def _open_budget_dialog(self, budget_btn=None, budget_text_ref=None, redraw_fn=None):
        """Modal to set/update the budget for this wallet month."""
        wallet = self._wallet
        root = self.winfo_toplevel()
        root.update_idletasks()
        x, y = root.winfo_rootx(), root.winfo_rooty()
        w, h = root.winfo_width(), root.winfo_height()
        try:
            from PIL import ImageGrab, ImageEnhance
            screenshot = ImageGrab.grab(bbox=(x, y, x + w, y + h))
            screenshot = ImageEnhance.Brightness(screenshot).enhance(0.5)
            _bg_photo = ImageTk.PhotoImage(screenshot)
        except Exception:
            _bg_photo = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_canvas = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_photo:
            ov_canvas.create_image(0, 0, anchor="nw", image=_bg_photo)
            ov_canvas._bg_photo = _bg_photo
        else:
            ov_canvas.config(bg="black")
            ov_canvas.create_rectangle(0, 0, w, h, fill="black", stipple="gray50", outline="")
        ov_canvas.bind("<Button-1>", lambda e: overlay.destroy())

        modal = tk.Frame(overlay, bg=BG_WHITE, padx=28, pady=24,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=420)
        modal.bind("<Button-1>", lambda e: "break")

        # header
        hdr = tk.Frame(modal, bg=BG_WHITE)
        hdr.pack(fill="x", pady=(0, 4))
        tk.Label(hdr, text="Add Budget", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(12, "bold")).pack(side="left")
        close_btn = tk.Label(hdr, text="\u00d7", bg=BG_WHITE, fg="#616161",
                             font=("Arial", 16), cursor="hand2")
        close_btn.pack(side="right")
        close_btn.bind("<Button-1>", lambda e: overlay.destroy())

        tk.Label(modal, text="Set or update the budget for this month.",
                 bg=BG_WHITE, fg="#616161", font=font(8)).pack(anchor="w", pady=(0, 14))

        # Budget amount field
        grp = tk.Frame(modal, bg=BG_WHITE)
        grp.pack(fill="x", pady=(0, 10))
        tk.Label(grp, text="Budget (PHP)", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(8, "bold")).pack(anchor="w")
        e_amount = tk.Entry(grp, font=font(9), bd=1, relief="solid",
                            highlightbackground=_FILTER_BDR, highlightthickness=1)
        e_amount.pack(fill="x", ipady=6, pady=(3, 0))
        # pre-fill existing budget
        try:
            existing = get_wallet_budget(wallet["id"], wallet.get("year"), wallet.get("month_id"))
            if existing:
                e_amount.insert(0, str(existing))
        except Exception:
            pass
        err_lbl = tk.Label(grp, text="", bg=BG_WHITE, fg="#C62828", font=font(7))
        err_lbl.pack(anchor="w")

        # footer
        footer = tk.Frame(modal, bg=BG_WHITE)
        footer.pack(fill="x", pady=(14, 0))

        cancel = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                          font=font(9, "bold"), padx=18, pady=8, cursor="hand2",
                          highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel.pack(side="left")
        cancel.bind("<Button-1>", lambda e: overlay.destroy())
        cancel.bind("<Enter>", lambda e: cancel.config(bg="#F5F1E8"))
        cancel.bind("<Leave>", lambda e: cancel.config(bg=BG_WHITE))

        save_btn = tk.Label(footer, text="Save", bg=_BTN_BROWN, fg="white",
                            font=font(9, "bold"), padx=18, pady=8, cursor="hand2")
        save_btn.pack(side="right")
        save_btn.bind("<Enter>", lambda e: save_btn.config(bg=_BTN_HOV))
        save_btn.bind("<Leave>", lambda e: save_btn.config(bg=_BTN_BROWN))

        def _save(e=None):
            val = e_amount.get().strip()
            if not val:
                err_lbl.config(text="Budget amount is required."); return
            try:
                amount = float(val)
            except ValueError:
                err_lbl.config(text="Must be a valid number."); return
            if amount < 0:
                err_lbl.config(text="Amount cannot be negative."); return
            try:
                upsert_wallet_budget(wallet["budget_id"], amount)
                if budget_text_ref is not None:
                    budget_text_ref[0] = f"Budget: Php {amount:,.2f}"
                if redraw_fn is not None:
                    redraw_fn()
                overlay.destroy()
            except Exception as err:
                err_lbl.config(text=f"Failed to save: {err}")

        save_btn.bind("<Button-1>", _save)

    def _open_report_dialog(self):
        wallet = self._wallet
        root   = self.winfo_toplevel()
        root.update_idletasks()
        x, y = root.winfo_rootx(), root.winfo_rooty()
        w, h = root.winfo_width(), root.winfo_height()
        try:
            from PIL import ImageGrab, ImageEnhance
            screenshot = ImageGrab.grab(bbox=(x, y, x + w, y + h))
            screenshot = ImageEnhance.Brightness(screenshot).enhance(0.5)
            _bg_photo  = ImageTk.PhotoImage(screenshot)
        except Exception:
            _bg_photo = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_canvas = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_photo:
            ov_canvas.create_image(0, 0, anchor="nw", image=_bg_photo)
            ov_canvas._bg_photo = _bg_photo
        else:
            ov_canvas.config(bg="black")
            ov_canvas.create_rectangle(0, 0, w, h, fill="black", stipple="gray50", outline="")
        ov_canvas.bind("<Button-1>", lambda e: overlay.destroy())

        modal = tk.Frame(overlay, bg=BG_WHITE,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=460, height=560)
        modal.bind("<Button-1>", lambda e: "break")

        # header
        hdr = tk.Frame(modal, bg=BG_WHITE)
        hdr.pack(fill="x", padx=28, pady=(18, 4))
        tk.Label(hdr, text="Generate Financial Report", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(12, "bold")).pack(side="left")
        close_btn = tk.Label(hdr, text="×", bg=BG_WHITE, fg="#616161",
                             font=("Arial", 16), cursor="hand2")
        close_btn.pack(side="right")
        close_btn.bind("<Button-1>", lambda e: overlay.destroy())

        tk.Label(modal, text="Fill in the details for this financial report.",
                 bg=BG_WHITE, fg="#616161", font=font(8)).pack(anchor="w", padx=28, pady=(0, 8))

        # scrollable body
        body_canvas = tk.Canvas(modal, bg=BG_WHITE, bd=0, highlightthickness=0)
        body_sb = ttk.Scrollbar(modal, orient="vertical", command=body_canvas.yview)
        body_inner = tk.Frame(body_canvas, bg=BG_WHITE)
        body_inner.bind("<Configure>",
                        lambda e: body_canvas.configure(scrollregion=body_canvas.bbox("all")))
        body_win = body_canvas.create_window((0, 0), window=body_inner, anchor="nw")
        body_canvas.configure(yscrollcommand=body_sb.set)
        body_canvas.bind("<Configure>",
                         lambda e: body_canvas.itemconfig(body_win, width=e.width))
        body_sb.pack(side="right", fill="y")
        body_canvas.pack(side="left", fill="both", expand=True)
        body_canvas.bind("<MouseWheel>",
                         lambda e: body_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        err_labels = {}

        def _field(label_text, key):
            grp = tk.Frame(body_inner, bg=BG_WHITE)
            grp.pack(fill="x", padx=28, pady=(0, 8))
            tk.Label(grp, text=label_text, bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(8, "bold")).pack(anchor="w")
            return grp

        def _entry(grp, key, readonly=False, prefill=""):
            e = tk.Entry(grp, font=font(9), bd=1, relief="solid",
                         highlightbackground=_FILTER_BDR, highlightthickness=1)
            e.pack(fill="x", ipady=6, pady=(3, 0))
            if prefill:
                e.insert(0, str(prefill))
            if readonly:
                e.config(state="readonly")
            err = tk.Label(grp, text="", bg=BG_WHITE, fg="#C62828", font=font(7))
            err.pack(anchor="w")
            err_labels[key] = err
            return e

        def _show_err(key, msg):
            if key in err_labels: err_labels[key].config(text=msg)
        def _clear_err(key):
            if key in err_labels: err_labels[key].config(text="")

        try:
            _budget  = get_wallet_budget(wallet["id"], wallet.get("year"), wallet.get("month_id"))
            _txs     = get_wallet_transactions(wallet["id"], budget_id=wallet.get("budget_id"))
            _income  = sum(t["price"]*t["quantity"] for t in _txs if t["kind"]=="income")
            _expense = sum(t["price"]*t["quantity"] for t in _txs if t["kind"]=="expense")
        except Exception:
            _budget = _income = _expense = 0.0

        try:
            _existing = get_wallet_reports(self._org.get("id"), wallet["id"])
            _rep_no   = f"FR-{len(_existing)+1:03d}"
        except Exception:
            _rep_no = "FR-001"

        e_event  = _entry(_field("Name of the event", "event"), "event")
        e_date   = _entry(_field("Date prepared", "date"), "date",
                          prefill=datetime.now().strftime("%Y-%m-%d"))
        e_repno  = _entry(_field("Financial report no.", "repno"), "repno",
                          readonly=True, prefill=_rep_no)
        e_bud    = _entry(_field("Budget for the month", "budget"), "budget",
                          prefill=_budget)
        e_inc    = _entry(_field(f"Total amount of income ({_income:,.2f})", "income"),
                          "income", prefill=_income)
        e_exp    = _entry(_field("Total amount of expenses", "expense"), "expense",
                          prefill=_expense)
        e_reimb  = _entry(_field("Reimbursement of expenses", "reimb"), "reimb", prefill="0")
        e_prev   = _entry(_field("Previous remaining fund", "prev"), "prev", prefill="0")
        e_bank   = _entry(_field("Budget in bank", "bank"), "bank", prefill="0")

        # footer inside body_inner
        footer = tk.Frame(body_inner, bg=BG_WHITE)
        footer.pack(fill="x", padx=28, pady=(10, 16))

        cancel = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                          font=font(9, "bold"), padx=18, pady=8, cursor="hand2",
                          highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel.pack(side="left")
        cancel.bind("<Button-1>", lambda e: overlay.destroy())
        cancel.bind("<Enter>", lambda e: cancel.config(bg="#F5F1E8"))
        cancel.bind("<Leave>", lambda e: cancel.config(bg=BG_WHITE))

        save_btn = tk.Label(footer, text="Continue", bg=_BTN_BROWN, fg="white",
                            font=font(9, "bold"), padx=18, pady=8, cursor="hand2")
        save_btn.pack(side="right")
        save_btn.bind("<Enter>", lambda e: save_btn.config(bg=_BTN_HOV))
        save_btn.bind("<Leave>", lambda e: save_btn.config(bg=_BTN_BROWN))

        def _save(e=None):
            valid = True
            vals = {
                "event":  e_event.get().strip(),
                "date":   e_date.get().strip(),
                "budget": e_bud.get().strip(),
                "income": e_inc.get().strip(),
                "expense":e_exp.get().strip(),
                "reimb":  e_reimb.get().strip(),
                "prev":   e_prev.get().strip(),
                "bank":   e_bank.get().strip(),
            }
            for key, val in vals.items():
                if not val:
                    _show_err(key, "This field is required."); valid = False
                else:
                    try:
                        if key not in ("event", "date"): float(val)
                        _clear_err(key)
                    except ValueError:
                        _show_err(key, "Must be a valid number."); valid = False
            if not valid:
                return
            # close form then open confirm dialog after event loop processes destroy
            _repno_val = e_repno.get().strip()
            overlay.destroy()
            self.after(50, lambda: self._open_confirm_generate(vals, _repno_val, wallet, getattr(self, "_report_banner_canvas", None)))

        save_btn.bind("<Button-1>", _save)


    def _open_confirm_generate(self, vals, report_no, wallet, banner_canvas=None):
        """Confirm dialog before saving the financial report to Supabase."""
        root = self.winfo_toplevel()
        root.update_idletasks()
        x, y = root.winfo_rootx(), root.winfo_rooty()
        w, h = root.winfo_width(), root.winfo_height()
        try:
            from PIL import ImageGrab, ImageEnhance
            screenshot = ImageGrab.grab(bbox=(x, y, x + w, y + h))
            screenshot = ImageEnhance.Brightness(screenshot).enhance(0.5)
            _bg_photo  = ImageTk.PhotoImage(screenshot)
        except Exception:
            _bg_photo = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_canvas = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_photo:
            ov_canvas.create_image(0, 0, anchor="nw", image=_bg_photo)
            ov_canvas._bg_photo = _bg_photo
        else:
            ov_canvas.config(bg="black")
            ov_canvas.create_rectangle(0, 0, w, h, fill="black", stipple="gray50", outline="")

        modal = tk.Frame(overlay, bg=BG_WHITE,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=400)
        modal.bind("<Button-1>", lambda e: "break")

        hdr = tk.Frame(modal, bg=BG_WHITE)
        hdr.pack(fill="x", padx=28, pady=(20, 4))
        tk.Label(hdr, text="Generate Report", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(12, "bold")).pack(side="left")
        close_x = tk.Label(hdr, text="×", bg=BG_WHITE, fg="#616161",
                           font=("Arial", 16), cursor="hand2")
        close_x.pack(side="right")
        close_x.bind("<Button-1>", lambda e: overlay.destroy())

        tk.Label(modal,
                 text="Generate the financial report for this wallet?\nThis may take a moment.",
                 bg=BG_WHITE, fg="#616161", font=font(9),
                 justify="left").pack(anchor="w", padx=28, pady=(0, 20))

        footer = tk.Frame(modal, bg=BG_WHITE)
        footer.pack(fill="x", padx=28, pady=(0, 20))

        cancel = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                          font=font(9, "bold"), padx=18, pady=8, cursor="hand2",
                          highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel.pack(side="left")
        cancel.bind("<Button-1>", lambda e: overlay.destroy())
        cancel.bind("<Enter>", lambda e: cancel.config(bg="#F5F1E8"))
        cancel.bind("<Leave>", lambda e: cancel.config(bg=BG_WHITE))

        yes_btn = tk.Label(footer, text="Yes, generate", bg=_BTN_BROWN, fg="white",
                           font=font(9, "bold"), padx=18, pady=8, cursor="hand2")
        yes_btn.pack(side="right")
        yes_btn.bind("<Enter>", lambda e: yes_btn.config(bg=_BTN_HOV))
        yes_btn.bind("<Leave>", lambda e: yes_btn.config(bg=_BTN_BROWN))

        def _confirm(e=None):
            try:
                add_financial_report(
                    org_id         = self._org.get("id"),
                    wallet_id      = wallet["id"],
                    budget_id      = wallet["budget_id"],
                    event_name     = vals["event"],
                    date_prepared  = vals["date"],
                    report_no      = report_no,
                    budget         = float(vals["budget"]),
                    total_income   = float(vals["income"]),
                    total_expense  = float(vals["expense"]),
                    reimbursement  = float(vals["reimb"]),
                    prev_fund      = float(vals["prev"]),
                    budget_in_bank = float(vals["bank"]),
                )
                overlay.destroy()
                if banner_canvas is not None:
                    self._show_report_action_buttons(banner_canvas)
                if self._active_tab == "archives":
                    self._switch_tab("archives")
            except Exception as err:
                tk.Label(modal, text=f"Error: {err}", bg=BG_WHITE,
                         fg="#C62828", font=font(8)).pack(padx=28)

        yes_btn.bind("<Button-1>", _confirm)

    def _show_report_action_buttons(self, banner_canvas, already_submitted=False):
        """Replace Generate button with Edit/Preview/Print/Submit buttons.
        If already_submitted=True, remove all four buttons and show a single
        hollow disabled 'Generate report' button — same pill as the original
        but greyed out and not clickable."""
        banner_canvas.delete("btn")
        # also remove any leftover individual action button windows
        for tag in ("btn_edit", "btn_preview", "btn_print", "btn_submit"):
            banner_canvas.delete(tag)

        _btn_h = 36

        if already_submitted:
            # ── Hollow disabled "Generate report" pill ────────────────
            _rep_ico = _img("reports.png", 14, 14, self._imgs)
            dis_cv = tk.Canvas(banner_canvas, bd=0, highlightthickness=0,
                               bg="#ECB95D", cursor="arrow")
            banner_canvas.create_window(0, 0, anchor="e", window=dis_cv, tags="btn")

            def _draw_disabled():
                dis_cv.delete("all")
                bw = dis_cv.winfo_width()
                bh = dis_cv.winfo_height()
                if bw < 4 or bh < 4:
                    return
                from PIL import Image, ImageDraw
                scale = 4
                sw, sh = bw * scale, bh * scale
                r = min(sh // 2, 20 * scale)
                # white fill at ~40% opacity blended over the gradient bg colour
                # get current bg colour of the canvas to blend against
                try:
                    bg_hex = dis_cv.cget("bg").lstrip("#")
                    br = int(bg_hex[0:2], 16)
                    bg_ = int(bg_hex[2:4], 16)
                    bb  = int(bg_hex[4:6], 16)
                except Exception:
                    br, bg_, bb = 0xEC, 0xB9, 0x5D   # fallback mid-gradient
                alpha = 0.65   # 65% white over gradient
                fr = int(255 * alpha + br * (1 - alpha))
                fg_ = int(255 * alpha + bg_ * (1 - alpha))
                fb  = int(255 * alpha + bb * (1 - alpha))
                img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
                ImageDraw.Draw(img).rounded_rectangle(
                    [0, 0, sw - 1, sh - 1], radius=r,
                    fill=(fr, fg_, fb, 255),
                    outline=(200, 190, 175, 255), width=8)
                img = img.resize((bw, bh), Image.LANCZOS)
                ph = ImageTk.PhotoImage(img)
                dis_cv._ph = ph
                dis_cv.create_image(0, 0, anchor="nw", image=ph)
                x = 12
                if _rep_ico:
                    dis_cv.create_image(x + 7, bh // 2, anchor="w",
                                        image=_rep_ico)
                    x += 26
                dis_cv.create_text(x, bh // 2, anchor="w",
                                   text="Generate report",
                                   fill="#AAAAAA", font=font(9, "bold"))

            _btn_w, _btn_h2 = 168, 36
            dis_cv.config(width=_btn_w, height=_btn_h2)
            dis_cv.bind("<Configure>", lambda e: _draw_disabled())

            def _layout_disabled(event=None):
                bw = banner_canvas.winfo_width()
                bh = banner_canvas.winfo_height()
                if bw < 2 or bh < 2:
                    return
                banner_canvas.coords("btn", bw - 20, bh // 2)
                # match gradient bg at button position
                c1, c2, c3 = (0xF3, 0xD5, 0x8D), (0xEC, 0xB9, 0x5D), (0xE5, 0x9E, 0x2C)
                t = max(0.0, min(1.0, (bw - _btn_w / 2) / bw)) if bw > 0 else 0.9
                if t <= 0.54:
                    s = t / 0.54
                    bc = tuple(int(c1[i] + (c2[i] - c1[i]) * s) for i in range(3))
                else:
                    s = (t - 0.54) / 0.46
                    bc = tuple(int(c2[i] + (c3[i] - c2[i]) * s) for i in range(3))
                dis_cv.config(bg="#{:02x}{:02x}{:02x}".format(*bc))

            banner_canvas.bind("<Configure>",
                lambda e, orig=banner_canvas.bind("<Configure>"): (
                    orig(e) if callable(orig) else None,
                    _layout_disabled()
                ))
            banner_canvas.after(50, _layout_disabled)
            return

        # ── Active: four action buttons ───────────────────────────────
        # pre-load icons at the same size as the Generate button icon (14×14)
        _ico_preview = _img("preview.png", 14, 14, self._imgs)
        _ico_print   = _img("print.png",   14, 14, self._imgs)
        _ico_submit  = _img("submit.png",  14, 14, self._imgs)

        def _make_action_btn(text, tag, icon=None,
                             fill_rgb=(255,255,255), hover_rgb=(236,220,198),
                             txt_clr=TEXT_DARK):
            cv = tk.Canvas(banner_canvas, bd=0, highlightthickness=0,
                           bg="#ECB95D", cursor="hand2",
                           width=90, height=_btn_h)
            banner_canvas.create_window(0, 0, anchor="e", window=cv, tags=tag)

            def _draw(hover=False):
                cv.delete("all")
                bw = cv.winfo_width() or 90
                bh = cv.winfo_height() or _btn_h
                from PIL import Image, ImageDraw
                scale = 4
                sw, sh = bw * scale, bh * scale
                r = min(sh // 2, 20 * scale)
                fill    = hover_rgb + (255,) if hover else fill_rgb + (255,)
                outline = (236, 220, 198, 255)
                img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
                ImageDraw.Draw(img).rounded_rectangle(
                    [0, 0, sw-1, sh-1], radius=r,
                    fill=fill, outline=outline, width=6)
                img = img.resize((bw, bh), Image.LANCZOS)
                ph = ImageTk.PhotoImage(img)
                cv._ph = ph
                cv.create_image(0, 0, anchor="nw", image=ph)
                # icon + text layout — same spacing as Generate button
                x = 12
                if icon:
                    cv.create_image(x + 7, bh // 2, anchor="w", image=icon)
                    x += 26
                cv.create_text(x, bh // 2, anchor="w",
                               text=text, fill=txt_clr, font=font(8, "bold"))

            cv.bind("<Configure>", lambda e: _draw())
            cv.bind("<Enter>",     lambda e: _draw(hover=True))
            cv.bind("<Leave>",     lambda e: _draw(hover=False))
            return cv

        # colours from wallets.css:
        #   preview → white / #ecddc6 hover  (default)
        #   print   → #a24a00 / #8b3a00 hover, white text
        #   submit  → #2e7d32 / #1b5e20 hover, white text
        edit_btn    = _make_action_btn("Edit Report", "btn_edit")
        preview_btn = _make_action_btn("Preview",     "btn_preview", icon=_ico_preview)
        print_btn   = _make_action_btn("Print",       "btn_print",   icon=_ico_print,
                                       fill_rgb=(162, 74, 0),
                                       hover_rgb=(139, 58, 0),
                                       txt_clr="white")
        submit_btn  = _make_action_btn("Submit",      "btn_submit",  icon=_ico_submit,
                                       fill_rgb=(46, 125, 50),
                                       hover_rgb=(27, 94, 32),
                                       txt_clr="white")

        edit_btn.bind("<Button-1>",    lambda e: self._open_report_dialog())
        preview_btn.bind("<Button-1>", lambda e: self._download_report_docx())
        print_btn.bind("<Button-1>",   lambda e: self._print_report())
        submit_btn.bind("<Button-1>",  lambda e: self._confirm_submit_report(banner_canvas))

        def _layout_action_btns(event=None):
            bw = banner_canvas.winfo_width()
            bh = banner_canvas.winfo_height()
            if bw < 2 or bh < 2:
                return
            cy  = bh // 2
            gap = 6
            bw_ = 90
            x = bw - 20
            for tag in ("btn_submit", "btn_print", "btn_preview", "btn_edit"):
                banner_canvas.coords(tag, x, cy)
                x -= bw_ + gap
            c1, c2, c3 = (0xF3, 0xD5, 0x8D), (0xEC, 0xB9, 0x5D), (0xE5, 0x9E, 0x2C)
            for cv_widget, x_pos in [
                (edit_btn,    bw - 20 - 3*(bw_+gap) - bw_//2),
                (preview_btn, bw - 20 - 2*(bw_+gap) - bw_//2),
                (print_btn,   bw - 20 - 1*(bw_+gap) - bw_//2),
                (submit_btn,  bw - 20 - bw_//2),
            ]:
                t = max(0.0, min(1.0, x_pos / bw)) if bw > 0 else 0.9
                if t <= 0.54:
                    s = t / 0.54
                    bc = tuple(int(c1[i] + (c2[i] - c1[i]) * s) for i in range(3))
                else:
                    s = (t - 0.54) / 0.46
                    bc = tuple(int(c2[i] + (c3[i] - c2[i]) * s) for i in range(3))
                cv_widget.config(bg="#{:02x}{:02x}{:02x}".format(*bc))

        banner_canvas.bind("<Configure>",
            lambda e, orig=banner_canvas.bind("<Configure>"): (
                orig(e) if callable(orig) else None,
                _layout_action_btns()
            ))
        banner_canvas.after(50, _layout_action_btns)

    def _download_report_docx(self):
        """Generate and save the financial report docx using the template."""
        import os as _os
        from docx import Document
        from tkinter import filedialog, messagebox
        from datetime import datetime as _dt

        wallet = self._wallet
        try:
            rep = get_latest_report(
                self._org.get("id"), wallet["id"], wallet["budget_id"])
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror("Error", f"Could not fetch report: {e}")
            return
        if not rep:
            from tkinter import messagebox
            messagebox.showwarning("No Report", "No report found for this wallet.")
            return

        template_path = _os.path.join(
            _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))),
            "assets", "template", "finance_report_template.docx")

        try:
            doc = Document(template_path)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror("Error", f"Could not open template: {e}")
            return

        def _replace(old, new):
            new = str(new) if new is not None else ""
            for p in doc.paragraphs:
                if old in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(old, new)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.text = run.text.replace(old, new)

        budget_val    = float(rep.get("budget") or 0)
        total_income  = float(rep.get("total_income") or 0)
        total_expense = float(rep.get("total_expense") or 0)
        reimb         = float(rep.get("reimbursement") or 0)
        prev_fund     = float(rep.get("previous_fund") or 0)
        bank          = float(rep.get("budget_in_the_bank") or 0)
        remaining     = budget_val - total_expense - reimb + prev_fund
        month_label   = wallet.get("month_name", "").capitalize()
        report_month  = f"{month_label} {wallet.get('year', '')}".upper()

        _replace("{{COLLEGE_NAME}}",      self._org.get("department", ""))
        _replace("{{ORG_NAME}}",          self._org.get("org_name", ""))
        _replace("{{EVENT_NAME}}",        rep.get("event_name") or "")
        _replace("{{REPORT_MONTH}}",      report_month)
        _replace("{{DATE_PREPARED}}",     str(rep.get("date_prepared") or ""))
        _replace("{{REPORT_NO}}",         rep.get("report_no") or "")
        _replace("{{BUDGET}}",            f"PHP {budget_val:,.2f}")
        _replace("{{TOTAL_INCOME}}",      f"PHP {total_income:,.2f}")
        _replace("{{TOTAL_EXPENSE}}",     f"PHP {total_expense:,.2f}")
        _replace("{{REIMBURSEMENT}}",     f"PHP {reimb:,.2f}")
        _replace("{{PREVIOUS_FUND}}",     f"PHP {prev_fund:,.2f}")
        _replace("{{BUDGET_IN_THE_BANK}}",f"PHP {bank:,.2f}")
        _replace("{{TOTAL_REMAINING}}",   f"PHP {remaining:,.2f}")

        # fill expense rows
        try:
            txs = get_wallet_transactions(wallet["id"], "expense",
                                          budget_id=wallet["budget_id"])
        except Exception:
            txs = []

        expenses_table = None
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            hdr = " ".join(c.text for c in table.rows[1].cells).upper()
            if "DATE ISSUED" in hdr and "PARTICULARS" in hdr:
                expenses_table = table
                break

        if expenses_table and txs:
            while len(expenses_table.rows) > 3:
                expenses_table._tbl.remove(expenses_table.rows[2]._tr)
            summary_row = expenses_table.rows[-1]
            for tx in txs:
                new_row = expenses_table.add_row()
                expenses_table._tbl.remove(new_row._tr)
                summary_row._tr.addprevious(new_row._tr)
                dc, qc, pc, desc_c, tc = new_row.cells
                dc.text    = str(tx.get("date_issued", ""))[:10]
                qc.text    = str(tx.get("quantity", ""))
                pc.text    = tx.get("particulars") or ""
                desc_c.text= tx.get("description") or ""
                tc.text    = f"PHP {float(tx.get('price',0))*int(tx.get('quantity',0)):,.2f}"
            summary_row.cells[-1].text = f"PHP {total_expense:,.2f}"

        # fill income rows
        try:
            incomes = get_wallet_transactions(wallet["id"], "income",
                                              budget_id=wallet["budget_id"])
        except Exception:
            incomes = []

        income_table = None
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            hdr = " ".join(c.text for c in table.rows[1].cells).upper()
            if "TYPE OF INCOME" in hdr and "DATE ISSUED" in hdr:
                income_table = table
                break

        if income_table and incomes:
            while len(income_table.rows) > 3:
                income_table._tbl.remove(income_table.rows[2]._tr)
            summary_row_inc = income_table.rows[-1]
            for tx in incomes:
                new_row = income_table.add_row()
                income_table._tbl.remove(new_row._tr)
                summary_row_inc._tr.addprevious(new_row._tr)
                dc, qc, tc, desc_c, pc = new_row.cells
                dc.text    = str(tx.get("date_issued", ""))[:10]
                qc.text    = str(tx.get("quantity", ""))
                tc.text    = tx.get("income_type") or ""
                desc_c.text= tx.get("description") or ""
                pc.text    = f"PHP {float(tx.get('price',0)):,.2f}"
            summary_row_inc.cells[-1].text = f"PHP {total_income:,.2f}"

        # ── Receipts appendix — same as web preview route ─────────────
        # Fetch receipts scoped to this wallet month (wallet_id + budget_id)
        try:
            receipts = get_wallet_receipts(
                wallet["id"], budget_id=wallet.get("budget_id"))
        except Exception:
            receipts = []

        if receipts:
            from docx.shared import Inches
            from io import BytesIO as _BytesIO

            # "APPENDIX: RECEIPTS" heading — matches web preview route
            title_p = doc.add_paragraph()
            title_run = title_p.add_run("APPENDIX: RECEIPTS")
            title_run.bold = True

            for r in receipts:
                file_path = r.get("file_url") or ""
                if not file_path:
                    continue

                # download image bytes from Supabase Storage
                try:
                    file_bytes = download_receipt_bytes(file_path)
                except Exception:
                    continue   # skip if file missing or not accessible

                # caption: "YYYY-MM-DD - Description"  (centered, same as web)
                cap_p = doc.add_paragraph()
                cap_p.alignment = 1
                cap_p.add_run(
                    f"{(r.get('receipt_date') or '')[:10]} - "
                    f"{r.get('description') or ''}"
                )

                # embed image at 3 inches wide (same as web preview route)
                try:
                    img_stream = _BytesIO(file_bytes)
                    doc.add_picture(img_stream, width=Inches(3))
                    doc.paragraphs[-1].alignment = 1   # center the image paragraph
                except Exception:
                    pass   # PDF or unsupported format — skip silently

                # small spacer between receipts
                spacer = doc.add_paragraph()
                spacer.alignment = 1

        # ask user where to save
        default_name = rep.get("report_no") or "financial_report"
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"{default_name}.docx",
            title="Save Financial Report"
        )
        if not save_path:
            return

        try:
            doc.save(save_path)
            messagebox.showinfo("Saved", f"Report saved to:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not save: {e}")

    def _print_report(self):
        """
        Build the same HTML as the web's print_report.html and open it in the
        default browser.  The page has onload="window.print()" so the OS print
        dialog fires automatically — identical behaviour to the web Print button.
        """
        import os as _os
        import tempfile
        import webbrowser
        from tkinter import messagebox

        wallet = self._wallet
        try:
            rep = get_latest_report(
                self._org.get("id"), wallet["id"], wallet["budget_id"])
        except Exception as e:
            messagebox.showerror("Error", f"Could not fetch report: {e}")
            return
        if not rep:
            messagebox.showwarning("No Report", "No report found for this wallet.")
            return

        # ── numeric fields ────────────────────────────────────────────
        budget_val    = float(rep.get("budget") or 0)
        total_income  = float(rep.get("total_income") or 0)
        total_expense = float(rep.get("total_expense") or 0)
        reimb         = float(rep.get("reimbursement") or 0)
        prev_fund     = float(rep.get("previous_fund") or 0)
        bank          = float(rep.get("budget_in_the_bank") or 0)
        remaining     = budget_val - total_expense - reimb + prev_fund
        month_label   = wallet.get("month_name", "").capitalize()
        report_month  = f"{month_label} {wallet.get('year', '')}".upper()
        org_name      = self._org.get("org_name", "")
        college_name  = self._org.get("department", "").upper()

        # ── transactions ──────────────────────────────────────────────
        try:
            txs     = get_wallet_transactions(wallet["id"], "expense",
                                              budget_id=wallet.get("budget_id"))
            incomes = get_wallet_transactions(wallet["id"], "income",
                                              budget_id=wallet.get("budget_id"))
        except Exception:
            txs = incomes = []

        # ── receipts with public URLs ─────────────────────────────────
        try:
            receipts_raw = get_wallet_receipts(
                wallet["id"], budget_id=wallet.get("budget_id"))
        except Exception:
            receipts_raw = []

        receipts = []
        for r in receipts_raw:
            fp = r.get("file_url") or ""
            url = ""
            if fp:
                try:
                    url = get_receipt_public_url(fp)
                except Exception:
                    url = ""
            receipts.append({
                "receipt_date": (r.get("receipt_date") or "")[:10],
                "description":  r.get("description") or "",
                "file_url":     url,
            })

        # ── load logos as base64 data URIs so they work in any temp path ──
        import base64

        def _img_b64(name):
            path = _os.path.join(
                _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))),
                "assets", "images", name)
            if not _os.path.exists(path):
                return ""
            with open(path, "rb") as f:
                data = base64.b64encode(f.read()).decode("ascii")
            ext = name.rsplit(".", 1)[-1].lower()
            mime = "image/png" if ext == "png" else f"image/{ext}"
            return f"data:{mime};base64,{data}"

        logo1 = _img_b64("1.png")
        logo2 = _img_b64("2.png")

        logo1_tag = (f'<img src="{logo1}" class="header-logo" alt="Logo 1" />'
                     if logo1 else "")
        logo2_tag = (f'<img src="{logo2}" class="header-logo" alt="Logo 2" />'
                     if logo2 else "")
        def _expense_rows():
            rows = ""
            for tx in txs:
                total = float(tx.get("price", 0)) * int(tx.get("quantity", 1))
                rows += f"""
                <tr>
                  <td class="text-center">{tx.get('date_issued','')[:10]}</td>
                  <td class="text-center">{tx.get('quantity','')}</td>
                  <td class="text-center">{tx.get('particulars') or ''}</td>
                  <td class="text-center">{tx.get('description') or ''}</td>
                  <td class="text-right">PHP {total:,.2f}</td>
                </tr>"""
            return rows

        def _income_rows():
            rows = ""
            for tx in incomes:
                rows += f"""
                <tr>
                  <td class="text-center">{tx.get('date_issued','')[:10]}</td>
                  <td class="text-center">{tx.get('quantity','')}</td>
                  <td class="text-center">{tx.get('income_type') or ''}</td>
                  <td class="text-center">{tx.get('description') or ''}</td>
                  <td class="text-right">PHP {float(tx.get('price',0)):,.2f}</td>
                </tr>"""
            return rows

        def _receipt_items():
            if not receipts:
                return ""
            items = ""
            for r in receipts:
                img_tag = (f'<img src="{r["file_url"]}" alt="Receipt" '
                           f'class="receipt-image" />'
                           if r["file_url"] else "")
                items += f"""
                <div class="receipt-item">
                  <div class="receipt-caption">
                    {r['receipt_date']} - {r['description']}
                  </div>
                  {img_tag}
                </div>"""
            return items

        appendix_section = ""
        if receipts:
            appendix_section = f"""
            <div class="page-break"></div>
            <div class="header-row">
              <div class="header-cell left">{logo1_tag}</div>
              <div class="header-cell center">
                <p class="header-center-line">Republic of the Philippines</p>
                <p class="header-center-line header-university">Laguna State Polytechnic University</p>
                <p class="header-center-line">Province of Laguna</p>
              </div>
              <div class="header-cell right">{logo2_tag}</div>
            </div>
            <div class="appendix-title">APPENDIX A - RECEIPTS</div>
            <div class="appendix-grid">
              {_receipt_items()}
            </div>"""

        # ── assemble full HTML (mirrors print_report.html exactly) ────
        html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8"/>
  <title>Activity Financial Statement</title>
  <style>
    :root {{ --font-main:"Times New Roman",serif; --blue-text:#0070c0; --gray-header:#d9d9d9; }}
    * {{ box-sizing:border-box; }}
    body {{ font-family:var(--font-main); font-size:11pt; color:#000; margin:0; padding:0; }}
    @page {{ size:A4; margin:20mm; }}
    .page {{ width:100%; }}
    .header-row {{ display:table; width:100%; margin-top:6mm; margin-bottom:4pt; }}
    .header-cell {{ display:table-cell; vertical-align:middle; }}
    .header-cell.left,.header-cell.right {{ width:20%; text-align:center; }}
    .header-cell.center {{ width:60%; text-align:center; font-size:11pt; line-height:1.15; }}
    .header-logo {{ width:70px; height:70px; object-fit:contain; }}
    .header-center-line {{ margin:0; }}
    .header-university {{ font-weight:normal; font-family:"Old English Text MT","Times New Roman",serif; font-size:14pt; }}
    .org-title {{ text-align:center; margin-top:6pt; margin-bottom:2pt; }}
    .org-title .dept-name {{ font-weight:bold; color:#000; display:block; }}
    .org-title .org-name {{ font-weight:bold; color:var(--blue-text); display:block; }}
    .statement-title {{ text-align:center; font-weight:bold; margin:4pt 0; background:#bfbfbf; padding:2px 0; }}
    table {{ width:100%; border-collapse:collapse; font-size:10pt; }}
    th,td {{ border:1px solid #000; padding:2px 4px; vertical-align:middle; }}
    th {{ background:var(--gray-header); text-align:center; font-weight:bold; }}
    .section-label-cell {{ width:30%; font-weight:bold; }}
    .section-value-cell {{ width:70%; text-align:center; }}
    .section-heading {{ background:var(--gray-header); font-weight:bold; text-align:center; }}
    .summary-table {{ margin-top:4pt; }}
    .summary-table td:first-child {{ width:55%; }}
    .summary-table td:last-child {{ width:45%; text-align:center; }}
    .summary-highlight td {{ background:#e7e6e6; font-weight:bold; }}
    .note {{ margin-top:6pt; font-size:9pt; }}
    .text-center {{ text-align:center; }}
    .text-right {{ text-align:center; }}
    .page-break {{ page-break-after:always; margin-top:20mm; }}
    .appendix-title {{ text-align:center; font-weight:bold; margin-top:10pt; margin-bottom:12pt; }}
    .appendix-grid {{ display:grid; grid-template-columns:repeat(2,1fr); gap:12pt; margin-top:8pt; }}
    .receipt-item {{ break-inside:avoid; page-break-inside:avoid; }}
    .receipt-caption {{ font-size:9pt; margin-bottom:4pt; }}
    .receipt-image {{ width:100%; height:auto; max-height:260px; object-fit:contain; border:1px solid #000; }}
    @media print {{ html,body {{ height:auto; overflow:visible; }} }}
  </style>
</head>
<body onload="window.print()">
<div class="page">

  <div class="header-row">
    <div class="header-cell left">{logo1_tag}</div>
    <div class="header-cell center">
      <p class="header-center-line">Republic of the Philippines</p>
      <p class="header-center-line header-university">Laguna State Polytechnic University</p>
      <p class="header-center-line">Province of Laguna</p>
    </div>
    <div class="header-cell right">{logo2_tag}</div>
  </div>

  <div class="org-title">
    <span class="dept-name">{college_name}</span>
    <span class="org-name">{org_name}</span>
  </div>

  <div class="statement-title">ACTIVITY FINANCIAL STATEMENT</div>

  <table style="margin-top:4pt">
    <tr><td class="section-label-cell">Name of Organization:</td>
        <td class="section-value-cell">{org_name}</td></tr>
    <tr><td class="section-label-cell">Name of the Event:</td>
        <td class="section-value-cell">{rep.get('event_name') or ''}</td></tr>
    <tr><td class="section-label-cell">Report for the Month of:</td>
        <td class="section-value-cell">{report_month}</td></tr>
    <tr><td class="section-label-cell">Date Prepared:</td>
        <td class="section-value-cell">{rep.get('date_prepared') or ''}</td></tr>
  </table>
  <table style="margin-top:2pt;margin-bottom:6pt">
    <tr><td class="section-label-cell">Financial Report No:</td>
        <td class="section-value-cell">{rep.get('report_no') or ''}</td></tr>
  </table>

  <table>
    <tr><td colspan="5" class="section-heading">EXPENSES</td></tr>
    <tr><th>Date Issued</th><th>Quantity</th><th>Particulars</th>
        <th>Description</th><th>Total</th></tr>
    {_expense_rows()}
    <tr class="summary-highlight">
      <td colspan="4" class="text-right"><strong>Total Amount of Expenses</strong></td>
      <td class="text-right"><strong>PHP {total_expense:,.2f}</strong></td>
    </tr>
  </table>

  <table style="margin-top:6pt">
    <tr><td colspan="5" class="section-heading">INCOME</td></tr>
    <tr><th>Date Issued</th><th>Quantity</th><th>Type of Income</th>
        <th>Description</th><th>Price</th></tr>
    {_income_rows()}
    <tr class="summary-highlight">
      <td colspan="4" class="text-right"><strong>Total Amount of Income</strong></td>
      <td class="text-right"><strong>PHP {total_income:,.2f}</strong></td>
    </tr>
  </table>

  <table class="summary-table" style="margin-top:6pt">
    <tr style="background:var(--gray-header);font-weight:bold">
      <td>Budget:</td><td>PHP {budget_val:,.2f}</td></tr>
    <tr><td>Total Amount of Expenses:</td><td>PHP {total_expense:,.2f}</td></tr>
    <tr><td>Total Amount of Income:</td><td>PHP {total_income:,.2f}</td></tr>
    <tr><td>Reimbursement of Expenses:</td><td>PHP {reimb:,.2f}</td></tr>
    <tr><td>Previous Remaining Fund:</td><td>PHP {prev_fund:,.2f}</td></tr>
  </table>
  <table class="summary-table">
    <tr style="background:var(--gray-header);font-weight:bold">
      <td>Total Remaining Balance:</td><td>PHP {remaining:,.2f}</td></tr>
    <tr><td>Budget in Bank:</td><td>PHP {bank:,.2f}</td></tr>
  </table>

  <p class="note">*NOTE: See APPENDIX A for receipts.</p>

  {appendix_section}

</div>
</body>
</html>"""

        # ── write to temp file and open in browser ────────────────────
        try:
            tmp = tempfile.NamedTemporaryFile(
                mode="w", suffix=".html", delete=False,
                encoding="utf-8", prefix="pockitrack_print_")
            tmp.write(html)
            tmp.close()
            webbrowser.open(f"file:///{tmp.name.replace(_os.sep, '/')}")
        except Exception as e:
            messagebox.showerror("Print Error", f"Could not open print preview: {e}")

    def _confirm_submit_report(self, banner_canvas):
        """Overlay confirm dialog before submitting the report to admin."""
        root = self.winfo_toplevel()
        root.update_idletasks()
        rx, ry = root.winfo_rootx(), root.winfo_rooty()
        rw, rh = root.winfo_width(), root.winfo_height()

        try:
            from PIL import ImageGrab, ImageEnhance
            shot = ImageGrab.grab(bbox=(rx, ry, rx + rw, ry + rh))
            shot = ImageEnhance.Brightness(shot).enhance(0.5)
            _bg_ph = ImageTk.PhotoImage(shot)
        except Exception:
            _bg_ph = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_cv = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_cv.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_ph:
            ov_cv.create_image(0, 0, anchor="nw", image=_bg_ph)
            ov_cv._bg_ph = _bg_ph
        else:
            ov_cv.config(bg="black")
            ov_cv.create_rectangle(0, 0, rw, rh, fill="black",
                                   stipple="gray50", outline="")
        ov_cv.bind("<Button-1>", lambda e: overlay.destroy())

        modal = tk.Frame(overlay, bg=BG_WHITE, padx=0, pady=0,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=400)
        modal.bind("<Button-1>", lambda e: "break")

        # header
        hdr = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=18)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Submit Report", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(13, "bold")).pack(side="left")
        close_lbl = tk.Label(hdr, text="×", bg=BG_WHITE, fg="#616161",
                             font=("Arial", 20), cursor="hand2", width=2)
        close_lbl.pack(side="right")
        close_lbl.bind("<Button-1>", lambda e: overlay.destroy())
        close_lbl.bind("<Enter>", lambda e: close_lbl.config(bg="#F0F0F0"))
        close_lbl.bind("<Leave>", lambda e: close_lbl.config(bg=BG_WHITE))
        tk.Frame(modal, bg=_FILTER_BDR, height=2).pack(fill="x")

        # body
        body = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=22)
        body.pack(fill="x")
        tk.Label(body,
                 text=("Submit this financial report to the admin?\n\n"
                       "Once submitted, the report will be locked and\n"
                       "sent to OSAS for review. This cannot be undone."),
                 bg=BG_WHITE, fg=TEXT_DARK, font=font(10),
                 wraplength=320, justify="left").pack(anchor="w")

        # footer
        tk.Frame(modal, bg=_FILTER_BDR, height=2).pack(fill="x")
        footer = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=16)
        footer.pack(fill="x")

        cancel_btn = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                              font=font(9, "bold"), padx=20, pady=8, cursor="hand2",
                              highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel_btn.pack(side="left")
        cancel_btn.bind("<Button-1>", lambda e: overlay.destroy())
        cancel_btn.bind("<Enter>", lambda e: cancel_btn.config(bg="#F5F1E8"))
        cancel_btn.bind("<Leave>", lambda e: cancel_btn.config(bg=BG_WHITE))

        def _do_submit(e=None):
            yes_btn.config(text="Submitting…", bg="#888888")
            overlay.update()
            try:
                submit_report(
                    org_id    = self._org.get("id"),
                    wallet_id = self._wallet["id"],
                    budget_id = self._wallet["budget_id"],
                )
                overlay.destroy()
                # redraw buttons in disabled/submitted state
                self._show_report_action_buttons(banner_canvas, already_submitted=True)
                # switch to Archives tab so user sees the new entry
                self._switch_tab("archives")
            except Exception as ex:
                yes_btn.config(text="Yes, Submit", bg=_BTN_BROWN)
                tk.Label(body, text=f"Submit failed: {str(ex)[:70]}",
                         bg=BG_WHITE, fg="#C62828", font=font(8)).pack(
                    anchor="w", pady=(8, 0))

        yes_btn = tk.Label(footer, text="Yes, Submit", bg=_BTN_BROWN, fg="white",
                           font=font(9, "bold"), padx=20, pady=8, cursor="hand2")
        yes_btn.pack(side="right")
        yes_btn.bind("<Button-1>", _do_submit)
        yes_btn.bind("<Enter>", lambda e: yes_btn.config(bg=_BTN_HOV))
        yes_btn.bind("<Leave>", lambda e: yes_btn.config(bg=_BTN_BROWN))

    def _delete_tx(self, t):
        try:
            delete_wallet_transaction(t["id"])
            if self._active_tab == "transactions":
                self._load_transactions()
        except Exception as err:
            print("Delete error:", err)


    def _switch_tab(self, key):
        self._active_tab = key
        for k, btn in self._tab_btns.items():
            if k == key:
                btn.config(fg=TEXT_DARK, font=font(10, "bold"))
                # underline effect via a small frame below
            else:
                btn.config(fg="#616161", font=font(10))

        for w in self._tab_content.winfo_children():
            w.destroy()

        if key == "transactions":
            self._tab_transactions()
        elif key == "reports":
            self._tab_reports()
        elif key == "receipts":
            self._tab_receipts()
        elif key == "archives":
            self._tab_archives()

    # ── TRANSACTIONS TAB ──────────────────────────────────────────────
    def _tab_transactions(self):
        p = self._tab_content

        # filter tabs — pill-shaped, same rounding as sidebar
        flt_row = tk.Frame(p, bg=BG_WHITE)
        flt_row.pack(pady=(0, 12))
        self._flt_btns = {}
        for key, lbl in [("all","All"),("income","Income"),("expense","Expense")]:
            cv = self._make_pill_btn(
                flt_row, lbl,
                bg_hex=BG_WHITE, hover_hex=_FILTER_BDR,
                fg=TEXT_DARK, width=80, height=34,
                font_spec=font(9),
                side="left", padx=5)
            cv.bind("<Button-1>", lambda e, k=key: self._set_tx_filter(k))
            self._flt_btns[key] = cv
        self._set_tx_filter(self._tx_filter, refresh=False)
        self._update_flt_styles()

        # list area
        self._tx_list = tk.Frame(p, bg=BG_WHITE)
        self._tx_list.pack(fill="both", expand=True)
        self._load_transactions()

    def _set_tx_filter(self, key, refresh=True):
        self._tx_filter = key
        self._update_flt_styles()
        if refresh:
            self._load_transactions()

    def _update_flt_styles(self):
        for k, cv in self._flt_btns.items():
            if k == self._tx_filter:
                cv._active_bg = _FILTER_ACT
                cv._active_fg = "white"
            else:
                cv._active_bg = None
                cv._active_fg = None
            # trigger redraw via the stored _redraw callable
            if hasattr(cv, "_redraw"):
                cv.after(0, cv._redraw)

    def _load_transactions(self):
        for w in self._tx_list.winfo_children():
            w.destroy()

        loading = tk.Label(self._tx_list, text="Loading...",
                           bg=BG_WHITE, fg=TEXT_MUTED, font=font(9))
        loading.pack(pady=10)
        self.update()

        try:
            rows = get_wallet_transactions(self._wallet["id"], self._tx_filter,
                                           budget_id=self._wallet.get("budget_id"))
        except Exception:
            rows = []
        loading.destroy()

        if not rows:
            # ── empty state with navi_history.png icon ────────────────
            frm = tk.Frame(self._tx_list, bg=BG_WHITE)
            frm.pack(expand=True, pady=40)
            ico = _img("navi_history.png", 48, 48, self._imgs)
            if ico:
                tk.Label(frm, image=ico, bg=BG_WHITE).pack(pady=(0, 10))
            tk.Label(frm, text="No transactions found.",
                     bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(11, "bold")).pack()
            tk.Label(frm, text="Try a different filter.",
                     bg=BG_WHITE, fg=TEXT_MUTED,
                     font=font(9)).pack(pady=(4, 0))
            return

        # ── scrollable list — scrollbar only appears when content overflows ──
        canvas = tk.Canvas(self._tx_list, bg=BG_WHITE, bd=0, highlightthickness=0)
        sb = ttk.Scrollbar(self._tx_list, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas, bg=BG_WHITE)
        win = canvas.create_window((0, 0), window=inner, anchor="nw")

        def _on_inner_configure(e=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
            # only show scrollbar when content is taller than the canvas
            content_h = inner.winfo_reqheight()
            canvas_h  = canvas.winfo_height()
            if content_h > canvas_h:
                sb.pack(side="right", fill="y")
            else:
                sb.pack_forget()

        inner.bind("<Configure>", _on_inner_configure)
        canvas.bind("<Configure>", lambda e: (
            canvas.itemconfig(win, width=e.width),
            _on_inner_configure()
        ))
        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side="left", fill="both", expand=True)
        self._bind_scroll(canvas)

        for t in rows:
            self._tx_item(inner, t)

    def _tx_item(self, parent, t):
        date  = t.get("date_issued", "")[:10]
        desc  = t.get("description") or "—"
        qty   = t.get("quantity", 1)
        price = t.get("price", 0)
        total = qty * price
        amt   = total if t["kind"] == "income" else -total
        qty_line = f"{price:,.2f} x {qty} ({total:,.2f})"

        card = tk.Frame(parent, bg=BG_WHITE,
                        highlightbackground=_FILTER_BDR,
                        highlightthickness=1, padx=14, pady=10)
        card.pack(fill="x", pady=4, padx=2)
        card.bind("<Enter>", lambda e: card.config(highlightbackground=_ACTIVE_TAB))
        card.bind("<Leave>", lambda e: card.config(highlightbackground=_FILTER_BDR))

        left = tk.Frame(card, bg=BG_WHITE)
        left.pack(side="left", fill="x", expand=True)

        if t["kind"] == "income":
            income_type = t.get("income_type") or "—"
            try:
                month_label = datetime.strptime(date, "%Y-%m-%d").strftime("%B").upper()
            except Exception:
                month_label = date
            tk.Label(left, text=month_label, bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(9, "bold"), anchor="w").pack(anchor="w")
            tk.Label(left, text=f"{qty_line}  ·  {income_type}  ·  {desc}",
                     bg=BG_WHITE, fg=TEXT_MUTED, font=font(8), anchor="w").pack(anchor="w")
            tk.Label(left, text=date, bg=BG_WHITE, fg="#999999",
                     font=font(7), anchor="w").pack(anchor="w")
        else:
            particulars = t.get("particulars") or "—"
            try:
                month_label = datetime.strptime(date, "%Y-%m-%d").strftime("%B").upper()
            except Exception:
                month_label = date
            tk.Label(left, text=month_label, bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(9, "bold"), anchor="w").pack(anchor="w")
            tk.Label(left, text=f"{qty_line}  ·  {particulars}  ·  {desc}",
                     bg=BG_WHITE, fg=TEXT_MUTED, font=font(8), anchor="w").pack(anchor="w")
            tk.Label(left, text=date, bg=BG_WHITE, fg="#999999",
                     font=font(7), anchor="w").pack(anchor="w")

        right = tk.Frame(card, bg=BG_WHITE)
        right.pack(side="right", anchor="center")

        color = _INCOME_CLR if amt >= 0 else _EXPENSE_CLR
        sign  = "+" if amt >= 0 else "-"
        tk.Label(right, text=f"{sign}₱{abs(amt):,.2f}",
                 bg=BG_WHITE, fg=color,
                 font=font(10, "bold")).pack(side="left", padx=(0, 8))

        # 3-dot menu button
        dots = tk.Label(right, text="...", bg=BG_WHITE, fg="#616161",
                        font=("Arial", 14), cursor="hand2", padx=4)
        dots.pack(side="left")

        dot_drop = tk.Frame(self._box, bg="white",
                            highlightbackground="#ddd", highlightthickness=1)
        dot_drop._open = False

        def _toggle_dot(e, dd=dot_drop, d=dots):
            if dd._open:
                dd.place_forget(); dd._open = False
            else:
                dd.update_idletasks()
                # position dropdown just below the dots button
                bx = d.winfo_rootx() - self._box.winfo_rootx()
                by = d.winfo_rooty() - self._box.winfo_rooty() + d.winfo_height()
                dd.place(in_=self._box, x=bx - 80, y=by)
                dd.lift()
                dd._open = True

        def _close_dot(dd=dot_drop):
            dd.place_forget(); dd._open = False

        edit_lbl = tk.Label(dot_drop, text="Edit", bg="white", fg=TEXT_DARK,
                            font=font(9), padx=16, pady=8, anchor="w", cursor="hand2")
        edit_lbl.pack(fill="x")
        edit_lbl.bind("<Enter>", lambda e: edit_lbl.config(bg="#F5F1E8"))
        edit_lbl.bind("<Leave>", lambda e: edit_lbl.config(bg="white"))
        edit_lbl.bind("<Button-1>", lambda e, tx=t, dd=dot_drop: (
            _close_dot(dd),
            self._open_tx_dialog(tx["kind"], existing=tx)
        ))

        del_lbl = tk.Label(dot_drop, text="Delete", bg="white", fg="#C62828",
                           font=font(9), padx=16, pady=8, anchor="w", cursor="hand2")
        del_lbl.pack(fill="x")
        del_lbl.bind("<Enter>", lambda e: del_lbl.config(bg="#FFF0F0"))
        del_lbl.bind("<Leave>", lambda e: del_lbl.config(bg="white"))
        del_lbl.bind("<Button-1>", lambda e, tx=t, dd=dot_drop: (
            _close_dot(dd),
            self._delete_tx(tx)
        ))

        dots.bind("<Button-1>", _toggle_dot)


    def _tab_reports(self):
        p = self._tab_content

        try:
            budget  = get_wallet_budget(self._wallet["id"],
                                        self._wallet.get("year"),
                                        self._wallet.get("month_id"))
            txs     = get_wallet_transactions(self._wallet["id"],
                                              budget_id=self._wallet.get("budget_id"))
            income  = sum(t["price"]*t["quantity"] for t in txs if t["kind"]=="income")
            expense = sum(t["price"]*t["quantity"] for t in txs if t["kind"]=="expense")
            ending  = income - expense
        except Exception:
            budget = income = expense = ending = 0.0

        # ── check if a report already exists for this specific wallet month ──
        # We match on budget_id (unique per wallet+month) to ensure we don't
        # show action buttons for a different month's report on the same wallet.
        current_budget_id = self._wallet.get("budget_id")
        try:
            existing_report = get_latest_report(
                self._org.get("id"),
                self._wallet["id"],
                current_budget_id,
            ) if current_budget_id else {}
        except Exception:
            existing_report = {}

        # double-check the returned report actually belongs to this month
        if existing_report and existing_report.get("budget_id") != current_budget_id:
            existing_report = {}

        # check if already submitted — determines button state
        already_submitted = (existing_report.get("status") == "Submitted"
                             if existing_report else False)

        # gradient report header
        banner_canvas = tk.Canvas(p, bg=BG_WHITE, bd=0, highlightthickness=0, height=90)
        self._report_banner_canvas = banner_canvas
        banner_canvas.pack(fill="x", pady=(0, 16))

        reports_img = _img("reports.png", 44, 44, self._imgs)

        def _draw_report_banner(event=None):
            bw = banner_canvas.winfo_width()
            bh = banner_canvas.winfo_height()
            if bw < 2 or bh < 2:
                return
            from PIL import Image, ImageDraw
            scale = 4
            sw, sh, r = bw * scale, bh * scale, 25 * scale
            c1, c2, c3 = (0xF3, 0xD5, 0x8D), (0xEC, 0xB9, 0x5D), (0xE5, 0x9E, 0x2C)
            grad = Image.new("RGB", (sw, 1))
            for x in range(sw):
                t = x / (sw - 1)
                if t <= 0.54:
                    s = t / 0.54
                    px = tuple(int(c1[i] + (c2[i] - c1[i]) * s) for i in range(3))
                else:
                    s = (t - 0.54) / 0.46
                    px = tuple(int(c2[i] + (c3[i] - c2[i]) * s) for i in range(3))
                grad.putpixel((x, 0), px)
            grad = grad.resize((sw, sh), Image.NEAREST)
            mask = Image.new("L", (sw, sh), 0)
            ImageDraw.Draw(mask).rounded_rectangle([0, 0, sw-1, sh-1], radius=r, fill=255)
            result = Image.new("RGBA", (sw, sh), (255, 255, 255, 0))
            result.paste(grad, mask=mask)
            result = result.resize((bw, bh), Image.LANCZOS)
            ph = ImageTk.PhotoImage(result)
            banner_canvas._ph = ph
            banner_canvas.delete("banner_bg")
            banner_canvas.create_image(0, 0, anchor="nw", image=ph, tags="banner_bg")
            banner_canvas.tag_lower("banner_bg")
            # reposition icon + text
            cy = bh // 2
            x = 20
            if reports_img:
                banner_canvas.coords("ico", x, cy)
                x += 58
            banner_canvas.coords("title_txt", x, cy - 10)
            banner_canvas.coords("sub_txt",   x, cy + 12)
            banner_canvas.coords("btn",       bw - 20, cy)

        # icon
        if reports_img:
            ico_lbl = tk.Label(banner_canvas, image=reports_img, bg="#ECB95D", bd=0)
            banner_canvas.create_window(20, 0, anchor="w", window=ico_lbl, tags="ico")
            banner_canvas._ico = ico_lbl

        banner_canvas.create_text(0, 0, anchor="w", tags="title_txt",
                                  text="Generate Financial Report",
                                  fill=TEXT_DARK, font=font(11, "bold"))
        banner_canvas.create_text(0, 0, anchor="w", tags="sub_txt",
                                  text="Create an Activity Financial Statement for this wallet.",
                                  fill="#5a3a00", font=font(8))

        banner_canvas.bind("<Configure>", _draw_report_banner)

        # ── show action buttons if report exists, otherwise Generate button ──
        if existing_report:
            # report already generated — show action buttons
            # pass already_submitted so buttons are greyed if already sent
            self.after(50, lambda s=already_submitted:
                       self._show_report_action_buttons(banner_canvas, already_submitted=s))
        else:
            # no report yet — show the Generate button
            _rep_ico = _img("reports.png", 14, 14, self._imgs)
            gen_btn_cv = tk.Canvas(banner_canvas, bd=0, highlightthickness=0,
                                   bg="#ECB95D", cursor="hand2")
            banner_canvas.create_window(0, 0, anchor="e", window=gen_btn_cv, tags="btn")

            def _draw_gen_btn(hover=False):
                gen_btn_cv.delete("all")
                bw = gen_btn_cv.winfo_width()
                bh = gen_btn_cv.winfo_height()
                if bw < 4 or bh < 4:
                    return
                from PIL import Image, ImageDraw
                scale = 4
                sw, sh = bw * scale, bh * scale
                r = min(sh // 2, 20 * scale)
                fill_col = (236, 220, 198, 255) if hover else (255, 255, 255, 255)
                img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
                ImageDraw.Draw(img).rounded_rectangle(
                    [0, 0, sw-1, sh-1], radius=r,
                    fill=fill_col,
                    outline=(236, 220, 198, 255), width=6)
                img = img.resize((bw, bh), Image.LANCZOS)
                ph = ImageTk.PhotoImage(img)
                gen_btn_cv._ph = ph
                gen_btn_cv.create_image(0, 0, anchor="nw", image=ph)
                x = 12
                if _rep_ico:
                    gen_btn_cv.create_image(x + 7, bh // 2, anchor="w", image=_rep_ico)
                    x += 26
                gen_btn_cv.create_text(x, bh // 2, anchor="w",
                                       text="Generate report",
                                       fill=TEXT_DARK, font=font(9, "bold"))

            gen_btn_cv.bind("<Configure>", lambda e: _draw_gen_btn())
            gen_btn_cv.bind("<Enter>",     lambda e: _draw_gen_btn(hover=True))
            gen_btn_cv.bind("<Leave>",     lambda e: _draw_gen_btn(hover=False))
            gen_btn_cv.bind("<Button-1>",  lambda e: self._open_report_dialog())

            _btn_w = 168
            _btn_h = 36
            gen_btn_cv.config(width=_btn_w, height=_btn_h)

            # update gen button bg colour when banner redraws
            def _draw_report_banner_with_btn(event=None):
                _draw_report_banner(event)
                bw = banner_canvas.winfo_width()
                c1, c2, c3 = (0xF3, 0xD5, 0x8D), (0xEC, 0xB9, 0x5D), (0xE5, 0x9E, 0x2C)
                t_btn = max(0.0, min(1.0, (bw - _btn_w / 2) / bw)) if bw > 0 else 0.9
                if t_btn <= 0.54:
                    s = t_btn / 0.54
                    _bc = tuple(int(c1[i] + (c2[i] - c1[i]) * s) for i in range(3))
                else:
                    s = (t_btn - 0.54) / 0.46
                    _bc = tuple(int(c2[i] + (c3[i] - c2[i]) * s) for i in range(3))
                gen_btn_cv.config(bg="#{:02x}{:02x}{:02x}".format(*_bc))

            banner_canvas.bind("<Configure>", _draw_report_banner_with_btn)

        # stat cards - white with rounded border
        stats_row = tk.Frame(p, bg=BG_WHITE)
        stats_row.pack(fill="x", pady=(4, 0))

        _stat_items = [
            ("Budget",                   f"₱{budget:,.2f}"),
            ("Total amount of income",   f"₱{income:,.2f}"),
            ("Total amount of expenses", f"₱{expense:,.2f}"),
            ("Ending Cash",              f"₱{ending:,.2f}"),
        ]
        _n = len(_stat_items)

        def _make_stat_card(parent, col_idx, lbl_text, val_text):
            parent.columnconfigure(col_idx, weight=1)
            cv = tk.Canvas(parent, bd=0, highlightthickness=0, bg=BG_WHITE, height=80)
            cv.grid(row=0, column=col_idx, padx=6, sticky="nsew")
            lbl_w = tk.Label(cv, text=lbl_text, bg=BG_WHITE, fg="#616161",
                             font=font(7), wraplength=120, justify="center")
            val_w = tk.Label(cv, text=val_text, bg=BG_WHITE, fg=TEXT_DARK,
                             font=font(10, "bold"))

            def _draw(event=None):
                from PIL import Image, ImageDraw
                cw = cv.winfo_width()
                ch = cv.winfo_height()
                if cw < 4 or ch < 4:
                    return
                scale = 4
                sw, sh, r = cw * scale, ch * scale, 15 * scale
                bdr = int(int(_FILTER_BDR.lstrip("#"), 16) >> 0)
                br = (bdr >> 16) & 255
                bg_ = (bdr >> 8) & 255
                bb  = bdr & 255
                img = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
                d = ImageDraw.Draw(img)
                d.rounded_rectangle([0, 0, sw-1, sh-1], radius=r,
                                    fill=(255, 255, 255, 255),
                                    outline=(br, bg_, bb, 255), width=8)
                img = img.resize((cw, ch), Image.LANCZOS)
                ph = ImageTk.PhotoImage(img)
                cv._ph = ph
                cv.delete("card_bg")
                cv.create_image(0, 0, anchor="nw", image=ph, tags="card_bg")
                cv.tag_lower("card_bg")
                lbl_w.place(relx=0.5, y=12, anchor="n")
                val_w.place(relx=0.5, rely=0.55, anchor="n")

            cv.bind("<Configure>", _draw)

        for _i, (_lbl, _val) in enumerate(_stat_items):
            _make_stat_card(stats_row, _i, _lbl, _val)



    # ── RECEIPTS TAB ──────────────────────────────────────────────────
    def _tab_receipts(self):
        p = self._tab_content

        loading = tk.Label(p, text="Loading receipts...",
                           bg=BG_WHITE, fg=TEXT_MUTED, font=font(9))
        loading.pack(pady=10)
        self.update()

        try:
            # scope to this wallet month via budget_id
            receipts = get_wallet_receipts(
                self._wallet["id"],
                budget_id=self._wallet.get("budget_id"),
            )
        except Exception:
            receipts = []
        loading.destroy()

        # scrollable grid
        canvas = tk.Canvas(p, bg=BG_WHITE, bd=0, highlightthickness=0)
        sb = ttk.Scrollbar(p, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas, bg=BG_WHITE)
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self._bind_scroll(canvas)

        self._receipts_inner = inner   # keep ref for refresh after upload

        if not receipts:
            self._empty(inner, "No receipts found.",
                        "Use '+ Add → Add receipt' to upload one.")
            return

        receipts_ico = _img("receipts.png", 36, 36, self._imgs)
        cols = 4
        for i, r in enumerate(receipts):
            col_i = i % cols
            row_i = i // cols
            inner.columnconfigure(col_i, weight=1)
            self._receipt_card(inner, row_i, col_i, r, receipts_ico)

    def _receipt_card(self, parent, row, col, r, icon):
        card = tk.Frame(parent, bg=BG_WHITE,
                        highlightbackground=_FILTER_BDR,
                        highlightthickness=2, padx=14, pady=14)
        card.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

        # icon area
        ico_frm = tk.Frame(card, bg="#F5F5F5", width=60, height=60)
        ico_frm.pack(pady=(0, 8))
        ico_frm.pack_propagate(False)
        if icon:
            tk.Label(ico_frm, image=icon, bg="#F5F5F5").place(
                relx=0.5, rely=0.5, anchor="center")

        desc = r.get("description") or "Receipt"
        date = (r.get("receipt_date") or "")[:10]
        tk.Label(card, text=desc, bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(8, "bold"), wraplength=130).pack()
        tk.Label(card, text=date, bg=BG_WHITE, fg=TEXT_MUTED,
                 font=font(7)).pack(pady=(2, 10))

        # ── three action buttons ──────────────────────────────────────
        btn_row = tk.Frame(card, bg=BG_WHITE)
        btn_row.pack()

        def _action_btn(parent, text, bg, hover_bg, command):
            b = tk.Label(parent, text=text, bg=bg, fg=TEXT_DARK,
                         font=font(7, "bold"), padx=8, pady=4, cursor="hand2")
            b.pack(side="left", padx=2)
            b.bind("<Button-1>", lambda e: command())
            b.bind("<Enter>",    lambda e, w=b, hb=hover_bg: w.config(bg=hb))
            b.bind("<Leave>",    lambda e, w=b, ob=bg:       w.config(bg=ob))
            return b

        file_path = r.get("file_url") or ""
        receipt_id = r.get("id")

        _action_btn(btn_row, "View",
                    bg="#E3F2FD", hover_bg="#BBDEFB",
                    command=lambda fp=file_path: self._view_receipt(fp))

        _action_btn(btn_row, "Download",
                    bg="#E8F5E9", hover_bg="#C8E6C9",
                    command=lambda fp=file_path: self._download_receipt(fp))

        _action_btn(btn_row, "Delete",
                    bg="#FFEBEE", hover_bg="#FFCDD2",
                    command=lambda rid=receipt_id: self._confirm_delete_receipt(rid))

    def _view_receipt(self, file_path):
        """Fetch public URL and open in default browser."""
        import threading, webbrowser
        if not file_path:
            return

        def _open():
            try:
                url = get_receipt_public_url(file_path)
                if url:
                    webbrowser.open(url)
            except Exception:
                pass

        threading.Thread(target=_open, daemon=True).start()

    def _download_receipt(self, file_path):
        """Fetch download URL and open it (browser triggers download)."""
        import threading, webbrowser
        if not file_path:
            return

        def _dl():
            try:
                url = get_receipt_download_url(file_path)
                if url:
                    webbrowser.open(url)
            except Exception:
                pass

        threading.Thread(target=_dl, daemon=True).start()

    def _confirm_delete_receipt(self, receipt_id):
        """Overlay confirm dialog before deleting a receipt."""
        if not receipt_id:
            return

        root = self.winfo_toplevel()
        root.update_idletasks()
        rx, ry = root.winfo_rootx(), root.winfo_rooty()
        rw, rh = root.winfo_width(), root.winfo_height()

        try:
            from PIL import ImageGrab, ImageEnhance
            shot = ImageGrab.grab(bbox=(rx, ry, rx + rw, ry + rh))
            shot = ImageEnhance.Brightness(shot).enhance(0.5)
            _bg_ph = ImageTk.PhotoImage(shot)
        except Exception:
            _bg_ph = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_cv = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_cv.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_ph:
            ov_cv.create_image(0, 0, anchor="nw", image=_bg_ph)
            ov_cv._bg_ph = _bg_ph
        else:
            ov_cv.config(bg="black")
            ov_cv.create_rectangle(0, 0, rw, rh, fill="black",
                                   stipple="gray50", outline="")
        ov_cv.bind("<Button-1>", lambda e: overlay.destroy())

        modal = tk.Frame(overlay, bg=BG_WHITE, padx=0, pady=0,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=380)
        modal.bind("<Button-1>", lambda e: "break")

        # header
        hdr = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=18)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Delete Receipt", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(13, "bold")).pack(side="left")
        close_lbl = tk.Label(hdr, text="×", bg=BG_WHITE, fg="#616161",
                             font=("Arial", 20), cursor="hand2", width=2)
        close_lbl.pack(side="right")
        close_lbl.bind("<Button-1>", lambda e: overlay.destroy())
        close_lbl.bind("<Enter>", lambda e: close_lbl.config(bg="#F0F0F0"))
        close_lbl.bind("<Leave>", lambda e: close_lbl.config(bg=BG_WHITE))
        tk.Frame(modal, bg=_FILTER_BDR, height=2).pack(fill="x")

        # body
        body = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=22)
        body.pack(fill="x")
        tk.Label(body,
                 text="Are you sure you want to delete this receipt?\nThis cannot be undone.",
                 bg=BG_WHITE, fg=TEXT_DARK, font=font(10),
                 wraplength=300, justify="left").pack(anchor="w")

        # footer
        tk.Frame(modal, bg=_FILTER_BDR, height=2).pack(fill="x")
        footer = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=16)
        footer.pack(fill="x")

        cancel_btn = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                              font=font(9, "bold"), padx=20, pady=8, cursor="hand2",
                              highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel_btn.pack(side="left")
        cancel_btn.bind("<Button-1>", lambda e: overlay.destroy())
        cancel_btn.bind("<Enter>", lambda e: cancel_btn.config(bg="#F5F1E8"))
        cancel_btn.bind("<Leave>", lambda e: cancel_btn.config(bg=BG_WHITE))

        def _do_delete(e=None):
            del_btn.config(text="Deleting…", bg="#888888")
            overlay.update()
            try:
                delete_receipt(receipt_id)
                overlay.destroy()
                # refresh receipts tab
                if self._active_tab == "receipts":
                    self._switch_tab("receipts")
            except Exception as ex:
                del_btn.config(text="Delete", bg="#C62828")
                tk.Label(body, text=f"Failed: {str(ex)[:60]}",
                         bg=BG_WHITE, fg="#C62828", font=font(8)).pack(
                    anchor="w", pady=(8, 0))

        del_btn = tk.Label(footer, text="Delete", bg="#C62828", fg="white",
                           font=font(9, "bold"), padx=20, pady=8, cursor="hand2")
        del_btn.pack(side="right")
        del_btn.bind("<Button-1>", _do_delete)
        del_btn.bind("<Enter>", lambda e: del_btn.config(bg="#B71C1C"))
        del_btn.bind("<Leave>", lambda e: del_btn.config(bg="#C62828"))

    def _open_receipt_dialog(self):
        """Overlay modal to upload a receipt for the current wallet month."""
        from tkinter import filedialog
        wallet = self._wallet

        root = self.winfo_toplevel()
        root.update_idletasks()
        rx, ry = root.winfo_rootx(), root.winfo_rooty()
        rw, rh = root.winfo_width(), root.winfo_height()

        try:
            from PIL import ImageGrab, ImageEnhance
            shot = ImageGrab.grab(bbox=(rx, ry, rx + rw, ry + rh))
            shot = ImageEnhance.Brightness(shot).enhance(0.5)
            _bg_ph = ImageTk.PhotoImage(shot)
        except Exception:
            _bg_ph = None

        overlay = tk.Frame(root, bg="black")
        overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        overlay.lift()

        ov_cv = tk.Canvas(overlay, bd=0, highlightthickness=0)
        ov_cv.place(relx=0, rely=0, relwidth=1, relheight=1)
        if _bg_ph:
            ov_cv.create_image(0, 0, anchor="nw", image=_bg_ph)
            ov_cv._bg_ph = _bg_ph
        else:
            ov_cv.config(bg="black")
            ov_cv.create_rectangle(0, 0, rw, rh, fill="black",
                                   stipple="gray50", outline="")
        ov_cv.bind("<Button-1>", lambda e: overlay.destroy())

        modal = tk.Frame(overlay, bg=BG_WHITE, padx=0, pady=0,
                         highlightbackground="#E0D4C0", highlightthickness=1)
        modal.place(relx=0.5, rely=0.5, anchor="center", width=440)
        modal.bind("<Button-1>", lambda e: "break")

        # ── header ────────────────────────────────────────────────────
        hdr = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=18)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Add Receipt", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(13, "bold")).pack(side="left")
        close_lbl = tk.Label(hdr, text="×", bg=BG_WHITE, fg="#616161",
                             font=("Arial", 20), cursor="hand2", width=2)
        close_lbl.pack(side="right")
        close_lbl.bind("<Button-1>", lambda e: overlay.destroy())
        close_lbl.bind("<Enter>", lambda e: close_lbl.config(bg="#F0F0F0"))
        close_lbl.bind("<Leave>", lambda e: close_lbl.config(bg=BG_WHITE))
        tk.Frame(modal, bg=_FILTER_BDR, height=2).pack(fill="x")

        # ── body ──────────────────────────────────────────────────────
        body = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=20)
        body.pack(fill="x")

        err_labels = {}

        def _lbl_entry(label_text, key, prefill=""):
            grp = tk.Frame(body, bg=BG_WHITE)
            grp.pack(fill="x", pady=(0, 10))
            tk.Label(grp, text=label_text, bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(8, "bold")).pack(anchor="w")
            e = tk.Entry(grp, font=font(10), bd=1, relief="solid",
                         bg="#F9F9F9", fg=TEXT_DARK,
                         highlightbackground=_FILTER_BDR, highlightthickness=1)
            e.insert(0, prefill)
            e.pack(fill="x", ipady=7, pady=(4, 0))
            err = tk.Label(grp, text="", bg=BG_WHITE, fg="#C62828", font=font(7))
            err.pack(anchor="w")
            err_labels[key] = err
            return e

        e_desc = _lbl_entry("Description", "desc")
        e_date = _lbl_entry("Receipt Date (YYYY-MM-DD)", "date",
                            prefill=datetime.now().strftime("%Y-%m-%d"))

        # file picker row
        file_grp = tk.Frame(body, bg=BG_WHITE)
        file_grp.pack(fill="x", pady=(0, 4))
        tk.Label(file_grp, text="File (image or PDF)", bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(8, "bold")).pack(anchor="w")

        file_row = tk.Frame(file_grp, bg=BG_WHITE)
        file_row.pack(fill="x", pady=(4, 0))

        file_var = tk.StringVar(value="")
        file_lbl = tk.Label(file_row, textvariable=file_var,
                            text="No file chosen", bg="#F9F9F9", fg=TEXT_MUTED,
                            font=font(9), anchor="w", padx=10,
                            highlightbackground=_FILTER_BDR, highlightthickness=1)
        file_lbl.pack(side="left", fill="x", expand=True, ipady=7)

        chosen_path = {"path": ""}

        def _pick_file():
            path = filedialog.askopenfilename(
                title="Choose receipt file",
                filetypes=[
                    ("Images & PDF", "*.png *.jpg *.jpeg *.gif *.webp *.pdf"),
                    ("All files", "*.*"),
                ])
            if path:
                chosen_path["path"] = path
                fname = _os.path.basename(path)
                file_var.set(fname)
                file_lbl.config(fg=TEXT_DARK)
                err_labels["file"].config(text="")

        browse_btn = tk.Label(file_row, text="Browse", bg=_BTN_BROWN, fg="white",
                              font=font(8, "bold"), padx=14, pady=7, cursor="hand2")
        browse_btn.pack(side="right", padx=(8, 0))
        browse_btn.bind("<Button-1>", lambda e: _pick_file())
        browse_btn.bind("<Enter>", lambda e: browse_btn.config(bg=_BTN_HOV))
        browse_btn.bind("<Leave>", lambda e: browse_btn.config(bg=_BTN_BROWN))

        err_file = tk.Label(file_grp, text="", bg=BG_WHITE, fg="#C62828", font=font(7))
        err_file.pack(anchor="w")
        err_labels["file"] = err_file

        # ── footer ────────────────────────────────────────────────────
        tk.Frame(modal, bg=_FILTER_BDR, height=2).pack(fill="x")
        footer = tk.Frame(modal, bg=BG_WHITE, padx=28, pady=16)
        footer.pack(fill="x")

        cancel_btn = tk.Label(footer, text="Cancel", bg=BG_WHITE, fg="#616161",
                              font=font(9, "bold"), padx=20, pady=8, cursor="hand2",
                              highlightbackground=_FILTER_BDR, highlightthickness=1)
        cancel_btn.pack(side="left")
        cancel_btn.bind("<Button-1>", lambda e: overlay.destroy())
        cancel_btn.bind("<Enter>", lambda e: cancel_btn.config(bg="#F5F1E8"))
        cancel_btn.bind("<Leave>", lambda e: cancel_btn.config(bg=BG_WHITE))

        def _show_err(key, msg):
            if key in err_labels:
                err_labels[key].config(text=msg)

        def _upload(e=None):
            desc_val = e_desc.get().strip()
            date_val = e_date.get().strip()
            path_val = chosen_path["path"]

            valid = True
            if not desc_val:
                _show_err("desc", "Description is required."); valid = False
            else:
                err_labels["desc"].config(text="")
            if not date_val:
                _show_err("date", "Date is required."); valid = False
            else:
                err_labels["date"].config(text="")
            if not path_val:
                _show_err("file", "Please choose a file."); valid = False

            if not valid:
                return

            save_btn.config(text="Uploading…", bg="#888888")
            overlay.update()

            try:
                with open(path_val, "rb") as f:
                    file_bytes = f.read()

                add_wallet_receipt(
                    wallet_id    = wallet["id"],
                    budget_id    = wallet["budget_id"],
                    org_id       = self._org.get("id"),
                    org_name     = self._org.get("org_name", "org"),
                    wallet_name  = wallet.get("wallet_name", "wallet"),
                    month_name   = wallet.get("month_name", "month"),
                    file_bytes   = file_bytes,
                    filename     = _os.path.basename(path_val),
                    description  = desc_val,
                    receipt_date = date_val,
                )
                overlay.destroy()
                # refresh receipts tab if currently visible
                if self._active_tab == "receipts":
                    self._switch_tab("receipts")
            except Exception as ex:
                save_btn.config(text="Upload", bg=_BTN_BROWN)
                _show_err("file", f"Upload failed: {str(ex)[:60]}")

        save_btn = tk.Label(footer, text="Upload", bg=_BTN_BROWN, fg="white",
                            font=font(9, "bold"), padx=20, pady=8, cursor="hand2")
        save_btn.pack(side="right")
        save_btn.bind("<Button-1>", _upload)
        save_btn.bind("<Enter>", lambda e: save_btn.config(bg=_BTN_HOV))
        save_btn.bind("<Leave>", lambda e: save_btn.config(bg=_BTN_BROWN))

    # ── ARCHIVES TAB ──────────────────────────────────────────────────
    def _tab_archives(self):
        p = self._tab_content
        loading = tk.Label(p, text="Loading reports...",
                           bg=BG_WHITE, fg=TEXT_MUTED, font=font(9))
        loading.pack(pady=10)
        self.update()

        try:
            # Query financial_report_archives scoped to this wallet month
            reports = get_wallet_archives(
                self._org.get("id"),
                self._wallet["id"],
                self._wallet["budget_id"],
            )
        except Exception:
            reports = []
        loading.destroy()

        if not reports:
            self._empty(p, "No archived reports.",
                        "Submit a financial report to see it here.")
            return

        canvas = tk.Canvas(p, bg=BG_WHITE, bd=0, highlightthickness=0)
        sb = ttk.Scrollbar(p, orient="vertical", command=canvas.yview)
        inner = tk.Frame(canvas, bg=BG_WHITE)
        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=inner, anchor="nw")
        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self._bind_scroll(canvas)

        cols = 3
        for i, rep in enumerate(reports):
            col_i = i % cols
            row_i = i // cols
            inner.columnconfigure(col_i, weight=1)
            self._archive_card(inner, row_i, col_i, rep)

    def _archive_card(self, parent, row, col, rep):
        card = tk.Frame(parent, bg=BG_WHITE,
                        highlightbackground=_FILTER_BDR,
                        highlightthickness=2, padx=16, pady=14)
        card.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

        rep_no = rep.get("report_no") or f"FR-{rep['id']:03d}"
        tk.Label(card, text=rep_no, bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(10, "bold")).pack()

        event   = rep.get("event_name")   or "—"
        date    = (rep.get("date_prepared") or "")[:10]
        budget  = float(rep.get("budget")        or 0)
        expense = float(rep.get("total_expense") or 0)
        reimb   = float(rep.get("reimbursement") or 0)
        prev    = float(rep.get("previous_fund") or 0)
        remaining = float(rep.get("remaining")   or 0)

        for lbl, val in [
            ("Event",       event),
            ("Date",        date),
            ("Budget",      f"₱{budget:,.2f}"),
            ("Expenses",    f"₱{expense:,.2f}"),
            ("Remaining",   f"₱{remaining:,.2f}"),
        ]:
            row_f = tk.Frame(card, bg=BG_WHITE)
            row_f.pack(fill="x", pady=1)
            tk.Label(row_f, text=lbl+":", bg=BG_WHITE, fg=TEXT_MUTED,
                     font=font(7), anchor="w", width=10).pack(side="left")
            tk.Label(row_f, text=str(val), bg=BG_WHITE, fg=TEXT_DARK,
                     font=font(7, "bold"), anchor="w").pack(side="left")

        # Download button — pill-shaped, same rounding as sidebar
        archive_id = rep.get("id")
        dl_btn = self._make_pill_btn(
            card, "Download",
            bg_hex=_BTN_BROWN, hover_hex=_BTN_HOV,
            fg="white", width=110, height=30,
            font_spec=font(8, "bold"),
            command=lambda aid=archive_id: self._download_archive_docx(aid))
        dl_btn.pack(pady=(8, 0))

    def _download_archive_docx(self, archive_id):
        """
        Generate a docx from the archive snapshot and save it locally.
        Mirrors web's GET /api/archives/<archive_id>/download.
        """
        import os as _os
        from docx import Document
        from docx.shared import Inches
        from io import BytesIO as _BytesIO
        from tkinter import filedialog, messagebox

        # fetch archive summary
        try:
            arch_res = (_sb_client().table("financial_report_archives")
                        .select("*").eq("id", archive_id).limit(1).execute())
            if not arch_res.data:
                messagebox.showwarning("Not Found", "Archive record not found.")
                return
            arch = arch_res.data[0]
        except Exception as e:
            messagebox.showerror("Error", f"Could not fetch archive: {e}")
            return

        wallet_id = arch["wallet_id"]
        budget_id = arch["budget_id"]
        org_id    = self._org.get("id")

        template_path = _os.path.join(
            _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))),
            "assets", "template", "finance_report_template.docx")
        try:
            doc = Document(template_path)
        except Exception as e:
            messagebox.showerror("Error", f"Could not open template: {e}")
            return

        def _replace(old, new):
            new = str(new) if new is not None else ""
            for p in doc.paragraphs:
                if old in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(old, new)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.text = run.text.replace(old, new)

        budget_val    = float(arch.get("budget")        or 0)
        total_expense = float(arch.get("total_expense") or 0)
        reimb         = float(arch.get("reimbursement") or 0)
        prev_fund     = float(arch.get("previous_fund") or 0)
        remaining     = float(arch.get("remaining")     or 0)
        month_label   = self._wallet.get("month_name", "").capitalize()
        report_month  = f"{month_label} {self._wallet.get('year', '')}".upper()

        _replace("{{COLLEGE_NAME}}",       self._org.get("department", ""))
        _replace("{{ORG_NAME}}",           self._org.get("org_name", ""))
        _replace("{{EVENT_NAME}}",         arch.get("event_name") or "")
        _replace("{{REPORT_MONTH}}",       report_month)
        _replace("{{DATE_PREPARED}}",      str(arch.get("date_prepared") or ""))
        _replace("{{REPORT_NO}}",          arch.get("report_no") or "")
        _replace("{{BUDGET}}",             f"PHP {budget_val:,.2f}")
        _replace("{{TOTAL_EXPENSE}}",      f"PHP {total_expense:,.2f}")
        _replace("{{REIMBURSEMENT}}",      f"PHP {reimb:,.2f}")
        _replace("{{PREVIOUS_FUND}}",      f"PHP {prev_fund:,.2f}")
        _replace("{{TOTAL_REMAINING}}",    f"PHP {remaining:,.2f}")

        # fetch archived transactions
        try:
            tx_res = (_sb_client().table("financial_report_archive_transactions")
                      .select("*").eq("archive_id", archive_id)
                      .order("date_issued").execute())
            txs     = [t for t in (tx_res.data or []) if t.get("kind") == "expense"]
            incomes = [t for t in (tx_res.data or []) if t.get("kind") == "income"]
        except Exception:
            txs = incomes = []

        total_income = sum(float(t.get("price", 0)) * int(t.get("quantity", 1))
                          for t in incomes)
        _replace("{{TOTAL_INCOME}}",       f"PHP {total_income:,.2f}")
        _replace("{{BUDGET_IN_THE_BANK}}", f"PHP {remaining:,.2f}")

        # fill expense table
        expenses_table = None
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            hdr = " ".join(c.text for c in table.rows[1].cells).upper()
            if "DATE ISSUED" in hdr and "PARTICULARS" in hdr:
                expenses_table = table
                break
        if expenses_table and txs:
            while len(expenses_table.rows) > 3:
                expenses_table._tbl.remove(expenses_table.rows[2]._tr)
            summary_row = expenses_table.rows[-1]
            for tx in txs:
                new_row = expenses_table.add_row()
                expenses_table._tbl.remove(new_row._tr)
                summary_row._tr.addprevious(new_row._tr)
                dc, qc, pc, desc_c, tc = new_row.cells
                dc.text     = str(tx.get("date_issued", ""))[:10]
                qc.text     = str(tx.get("quantity", ""))
                pc.text     = tx.get("particulars") or ""
                desc_c.text = tx.get("description") or ""
                tc.text     = f"PHP {float(tx.get('price',0))*int(tx.get('quantity',1)):,.2f}"
            summary_row.cells[-1].text = f"PHP {total_expense:,.2f}"

        # fill income table
        income_table = None
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            hdr = " ".join(c.text for c in table.rows[1].cells).upper()
            if "TYPE OF INCOME" in hdr and "DATE ISSUED" in hdr:
                income_table = table
                break
        if income_table and incomes:
            while len(income_table.rows) > 3:
                income_table._tbl.remove(income_table.rows[2]._tr)
            summary_row_inc = income_table.rows[-1]
            for tx in incomes:
                new_row = income_table.add_row()
                income_table._tbl.remove(new_row._tr)
                summary_row_inc._tr.addprevious(new_row._tr)
                dc, qc, tc, desc_c, pc = new_row.cells
                dc.text     = str(tx.get("date_issued", ""))[:10]
                qc.text     = str(tx.get("quantity", ""))
                tc.text     = tx.get("income_type") or ""
                desc_c.text = tx.get("description") or ""
                pc.text     = f"PHP {float(tx.get('price',0)):,.2f}"
            summary_row_inc.cells[-1].text = f"PHP {total_income:,.2f}"

        # fetch archived receipts and embed as images
        try:
            rc_res = (_sb_client().table("financial_report_archive_receipts")
                      .select("*").eq("archive_id", archive_id)
                      .order("receipt_date").execute())
            arch_receipts = rc_res.data or []
        except Exception:
            arch_receipts = []

        if arch_receipts:
            title_p = doc.add_paragraph()
            title_p.add_run("APPENDIX: RECEIPTS").bold = True
            for r in arch_receipts:
                fp = r.get("file_url") or ""
                if not fp:
                    continue
                try:
                    file_bytes = download_receipt_bytes(fp)
                except Exception:
                    continue
                cap_p = doc.add_paragraph()
                cap_p.alignment = 1
                cap_p.add_run(f"{(r.get('receipt_date') or '')[:10]} - "
                              f"{r.get('description') or ''}")
                try:
                    doc.add_picture(_BytesIO(file_bytes), width=Inches(3))
                    doc.paragraphs[-1].alignment = 1
                except Exception:
                    pass
                spacer = doc.add_paragraph()
                spacer.alignment = 1

        default_name = arch.get("report_no") or "financial_report_archive"
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"{default_name}.docx",
            title="Save Archive Report")
        if not save_path:
            return
        try:
            doc.save(save_path)
            messagebox.showinfo("Saved", f"Report saved to:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not save: {e}")

    # ── helpers ───────────────────────────────────────────────────────
    def _clear_box(self):
        for w in self._box.winfo_children():
            w.destroy()

    def _empty(self, parent, title, subtitle):
        frm = tk.Frame(parent, bg=BG_WHITE)
        frm.pack(expand=True, pady=40)
        tk.Label(frm, text=title, bg=BG_WHITE, fg=TEXT_DARK,
                 font=font(11, "bold")).pack()
        tk.Label(frm, text=subtitle, bg=BG_WHITE, fg=TEXT_MUTED,
                 font=font(9)).pack(pady=(4, 0))

    def _bind_scroll(self, canvas):
        def _scroll(e):
            try:
                canvas.yview_scroll(int(-1*(e.delta/120)), "units")
            except Exception:
                pass
        canvas.bind_all("<MouseWheel>", _scroll)
        canvas.bind("<Destroy>",
                    lambda e: canvas.unbind_all("<MouseWheel>")
                    if e.widget is canvas else None)
